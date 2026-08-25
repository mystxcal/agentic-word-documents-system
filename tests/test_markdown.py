from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document
from docx.oxml.ns import qn

from agentic_docs.diagnostics import DiagnosticBag
from agentic_docs.errors import PackageError
from agentic_docs.jsonc import load_jsonc
from agentic_docs.sources.markdown import (
    compile_markdown,
    markdown_block_tag,
    markdown_component_tag,
    markdown_slot_tag,
    parse_blocks,
    parse_document,
)
from agentic_docs.word.fragments import component_wrapper, import_word_fragment, replace_nested_slot
from agentic_docs.word.ooxml import find_sdts
from agentic_docs.word.package import DOCUMENT_PART, DocxPackage, XNS, validate_docx


SYSTEM_ROOT = Path(__file__).resolve().parents[1]
STYLE_DONOR = SYSTEM_ROOT / "kits" / "studio" / "donors" / "Style_Gallery.docx"
KIT = load_jsonc(SYSTEM_ROOT / "kits" / "studio" / "kit.jsonc")


class MarkdownAdapterTests(unittest.TestCase):
    def test_blocks_keep_human_order_and_table_position(self):
        diagnostics = DiagnosticBag()
        blocks = parse_blocks(
            "# Heading\n\nBefore.\n\n| A | B |\n|---|---|\n| 1 | 2 |\n\nAfter.",
            diagnostics,
            "test.md",
        )
        self.assertEqual([block.kind for block in blocks], ["heading", "paragraph", "table", "paragraph"])
        self.assertFalse(diagnostics.values)

    def test_unknown_directive_is_visible_in_compatibility_mode_but_fails_strict_mode(self):
        diagnostics = DiagnosticBag()
        blocks = parse_blocks(":::unfamiliar value", diagnostics, "test.md")
        self.assertEqual(blocks[0].kind, "paragraph")
        self.assertEqual(blocks[0].value, ":::unfamiliar value")
        self.assertEqual(diagnostics.values[0].code, "MARKDOWN_UNKNOWN_DIRECTIVE")
        with self.assertRaisesRegex(PackageError, "Unsupported Markdown"):
            parse_document(":::unfamiliar value", DiagnosticBag(), "test.md", strict=True)

    def test_parses_nested_lists_callouts_hard_breaks_and_insertions(self):
        parsed = parse_document(
            "# Method {#method}\n\n"
            "First line.  \nSecond line.\n\n"
            "1. Parent\n  - Child\n\n"
            ":::note\nImportant **note**.\n:::\n\n"
            ":::insert reference-table",
            DiagnosticBag(),
            "method.md",
        )
        self.assertEqual(
            [block.kind for block in parsed.blocks],
            ["heading", "paragraph", "list", "callout", "insert"],
        )
        self.assertEqual(parsed.blocks[0].block_id, "method")
        self.assertIn("\n", parsed.blocks[1].value)
        self.assertEqual(parsed.blocks[2].value["items"][1]["level"], 1)
        self.assertEqual(parsed.slot_names, ["reference-table"])

    def test_markdown_image_requires_a_declared_figure_component(self):
        with self.assertRaisesRegex(PackageError, "declare a figure component"):
            parse_document(
                "![route sketch](route.png)",
                DiagnosticBag(),
                "method.md",
            )

    def test_compiles_semantic_word_content_agdoc_table_and_source_map(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "method.md"
            output = root / "method.docx"
            source.write_text(
                "# Observation method {#observation-method}\n\n"
                "Use **consistent notes** and refer to "
                "[the manufacturer](https://example.com/guide).\n\n"
                "| Item | Quantity |\n"
                "|:---|---:|\n"
                "| Shaded bench | 2 |\n\n"
                ":::insert reference-table",
                encoding="utf-8",
            )
            report = compile_markdown(
                source,
                output,
                component_id="method",
                style_source=STYLE_DONOR,
                semantic_styles=KIT["semantic_styles"],
                table_styles=KIT["table_styles"],
                declared_slots={"reference-table": ["data"]},
                options={"table_style_role": "technical"},
                total_width_twips=9000,
                diagnostics=DiagnosticBag(),
            )

            self.assertTrue(validate_docx(output)["valid"])
            self.assertEqual(report["slots"], ["reference-table"])
            self.assertEqual(report["source_map"][0]["block_id"], "observation-method")
            package = DocxPackage(output)
            document = package.xml(DOCUMENT_PART)
            self.assertEqual(len(find_sdts(document, markdown_component_tag("method"))), 1)
            heading_control = find_sdts(document, markdown_block_tag("method", "observation-method"))
            self.assertEqual(len(heading_control), 1)
            self.assertEqual(heading_control[0].getparent().tag, qn("w:p"))
            slot = find_sdts(document, markdown_slot_tag("method", "reference-table"))
            self.assertEqual(len(slot), 1)
            self.assertEqual(
                document.xpath("string(.//w:tbl[1]/w:tr[1]/w:tc[1]/w:tcPr/w:shd/@w:fill)", namespaces=XNS),
                "17324D",
            )
            _, relationships = package.relationship_root(DOCUMENT_PART)
            self.assertTrue(
                relationships.xpath(
                    "boolean(./pr:Relationship[contains(@Type, 'hyperlink') and @Target='https://example.com/guide'])",
                    namespaces=XNS,
                )
            )

    def test_list_levels_are_native_and_separate_lists_restart(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "method.md"
            output = root / "method.docx"
            source.write_text(
                "1. First\n"
                "  - Nested bullet\n"
                "2. Second\n\n"
                "Between lists.\n\n"
                "1. Restarted\n"
                "2. Again",
                encoding="utf-8",
            )
            compile_markdown(
                source,
                output,
                component_id="method",
                style_source=STYLE_DONOR,
                semantic_styles=KIT["semantic_styles"],
                table_styles=KIT["table_styles"],
                declared_slots={},
                options={},
                total_width_twips=9000,
                diagnostics=DiagnosticBag(),
            )
            package = DocxPackage(output)
            document = package.xml(DOCUMENT_PART)
            paragraphs = document.xpath(".//w:p[w:pPr/w:numPr]", namespaces=XNS)
            levels = [
                paragraph.xpath("string(w:pPr/w:numPr/w:ilvl/@w:val)", namespaces=XNS)
                for paragraph in paragraphs
            ]
            number_ids = [
                paragraph.xpath("string(w:pPr/w:numPr/w:numId/@w:val)", namespaces=XNS)
                for paragraph in paragraphs
            ]
            self.assertEqual(levels, ["0", "1", "0", "0", "0"])
            self.assertEqual(number_ids[0], number_ids[2])
            self.assertNotEqual(number_ids[0], number_ids[3])

    def test_markdown_insert_slot_accepts_an_independently_compiled_child(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "method.md"
            compiled = root / "method.docx"
            target_path = root / "target.docx"
            source.write_text(
                "Before insertion.\n\n:::insert reference-table\n\nAfter insertion.",
                encoding="utf-8",
            )
            compile_markdown(
                source,
                compiled,
                component_id="method",
                style_source=STYLE_DONOR,
                semantic_styles=KIT["semantic_styles"],
                table_styles=KIT["table_styles"],
                declared_slots={"reference-table": ["data"]},
                options={},
                total_width_twips=9000,
                diagnostics=DiagnosticBag(),
            )
            Document().save(target_path)
            target = DocxPackage(target_path)
            wrapper = import_word_fragment(
                target,
                compiled,
                component_id="method",
                source_tag=None,
                allow_untagged=True,
                preserve_sections=False,
                diagnostics=DiagnosticBag(),
            )
            child_document = Document()
            child_paragraph = child_document.add_paragraph("Excel-owned table goes here")._p
            replace_nested_slot(
                wrapper,
                markdown_slot_tag("method", "reference-table"),
                [component_wrapper("data", [child_paragraph])],
            )
            text = wrapper.xpath(".//w:t/text()", namespaces=XNS)
            self.assertEqual(
                text,
                ["Before insertion.", "Excel-owned table goes here", "After insertion."],
            )

    def test_declared_and_authored_slots_must_match(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "method.md"
            source.write_text(":::insert wrong-slot", encoding="utf-8")
            with self.assertRaisesRegex(PackageError, "undeclared"):
                compile_markdown(
                    source,
                    root / "out.docx",
                    component_id="method",
                    style_source=STYLE_DONOR,
                    semantic_styles=KIT["semantic_styles"],
                    table_styles=KIT["table_styles"],
                    declared_slots={"reference-table": ["data"]},
                    options={},
                    total_width_twips=9000,
                    diagnostics=DiagnosticBag(),
                )


if __name__ == "__main__":
    unittest.main()
