from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document
from lxml import etree

from agentic_docs.word.ooxml import wrap_elements
from agentic_docs.word.package import DOCUMENT_PART, DocxPackage, XNS
from agentic_docs.word.regions import apply_region_layout


class PageRegionTests(unittest.TestCase):
    def test_boundary_after_nested_cover_is_a_direct_body_paragraph(self):
        with TemporaryDirectory() as directory:
            path = Path(directory) / "nested-cover-boundary.docx"
            document = Document()
            cover = document.add_paragraph("Cover content")
            wrap_elements("AGDOC.COVER", [cover._p])
            main = document.add_paragraph("Main region")
            wrap_elements("AGDOC.REGION.MAIN", [main._p])
            document.save(path)

            package = DocxPackage(path)
            apply_region_layout(
                package,
                [
                    {"id": "front", "config": {"numbering": None}},
                    {
                        "id": "main",
                        "start_tag": "AGDOC.REGION.MAIN",
                        "boundary": "next_page",
                        "config": {"numbering": {"style": "arabic", "start": 1}},
                    },
                ],
            )

            root = package.xml(DOCUMENT_PART)
            body = root.find("w:body", namespaces=XNS)
            boundary = root.xpath(".//w:sectPr", namespaces=XNS)[0].getparent().getparent()
            self.assertEqual("p", etree.QName(boundary).localname)
            self.assertIs(boundary.getparent(), body)
            self.assertFalse(root.xpath(".//w:sdt//w:sectPr", namespaces=XNS))

    def test_three_regions_receive_independent_boundaries_margins_and_numbering(self):
        with TemporaryDirectory() as directory:
            path = Path(directory) / "regions.docx"
            document = Document()
            document.add_paragraph("Front matter")
            main = document.add_paragraph("Main region")
            wrap_elements("AGDOC.REGION.MAIN", [main._p])
            annex = document.add_paragraph("Annex region")
            wrap_elements("AGDOC.REGION.ANNEX", [annex._p])
            document.save(path)

            package = DocxPackage(path)
            result = apply_region_layout(
                package,
                [
                    {
                        "id": "front",
                        "config": {"top_margin_twips": 0, "numbering": None},
                    },
                    {
                        "id": "main",
                        "start_tag": "AGDOC.REGION.MAIN",
                        "boundary": "next_page",
                        "config": {
                            "top_margin_twips": 1800,
                            "bottom_margin_twips": 1200,
                            "left_margin_twips": 1300,
                            "right_margin_twips": 1400,
                            "numbering": {"style": "roman_lower", "start": 1, "page_count_scope": "region"},
                        },
                    },
                    {
                        "id": "annex",
                        "start_tag": "AGDOC.REGION.ANNEX",
                        "boundary": "continuous",
                        "config": {
                            "numbering": {"style": "arabic", "start": 5, "page_count_scope": "document"}
                        },
                    },
                ],
            )
            self.assertEqual(result["region_sections"], {"front": [1], "main": [2], "annex": [3]})
            self.assertEqual([item["boundary"] for item in result["boundaries"]], ["next_page", "continuous"])

            root = package.xml(DOCUMENT_PART)
            sections = root.xpath(".//w:sectPr", namespaces=XNS)
            main_margins = sections[1].find("w:pgMar", namespaces=XNS)
            self.assertEqual(main_margins.get(f"{{{XNS['w']}}}top"), "1800")
            self.assertEqual(main_margins.get(f"{{{XNS['w']}}}bottom"), "1200")
            self.assertEqual(main_margins.get(f"{{{XNS['w']}}}left"), "1300")
            self.assertEqual(main_margins.get(f"{{{XNS['w']}}}right"), "1400")
            main_numbering = sections[1].find("w:pgNumType", namespaces=XNS)
            annex_numbering = sections[2].find("w:pgNumType", namespaces=XNS)
            self.assertEqual(main_numbering.get(f"{{{XNS['w']}}}fmt"), "lowerRoman")
            self.assertEqual(main_numbering.get(f"{{{XNS['w']}}}start"), "1")
            self.assertEqual(annex_numbering.get(f"{{{XNS['w']}}}fmt"), "decimal")
            self.assertEqual(annex_numbering.get(f"{{{XNS['w']}}}start"), "5")

    def test_preserve_layout_mode_keeps_existing_section_numbering(self):
        with TemporaryDirectory() as directory:
            path = Path(directory) / "preserved.docx"
            Document().save(path)
            package = DocxPackage(path)
            root = package.xml(DOCUMENT_PART)
            section = root.xpath(".//w:sectPr", namespaces=XNS)[0]
            numbering = etree.Element(f"{{{XNS['w']}}}pgNumType")
            numbering.set(f"{{{XNS['w']}}}fmt", "lowerRoman")
            numbering.set(f"{{{XNS['w']}}}start", "7")
            section.append(numbering)
            package.set_xml(DOCUMENT_PART, root)

            result = apply_region_layout(
                package,
                [{"id": "main", "config": {"layout_mode": "preserve", "numbering": None}}],
            )

            section = package.xml(DOCUMENT_PART).xpath(".//w:sectPr", namespaces=XNS)[0]
            retained = section.find("w:pgNumType", namespaces=XNS)
            self.assertIsNotNone(retained)
            self.assertEqual(retained.get(f"{{{XNS['w']}}}fmt"), "lowerRoman")
            self.assertEqual(retained.get(f"{{{XNS['w']}}}start"), "7")
            self.assertEqual(result["layout_modes"], {"main": "preserve"})


if __name__ == "__main__":
    unittest.main()
