from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document
from PIL import Image

from agentic_docs.build_comparison import compare_build_candidate, comparison_page_furniture


class BuildComparisonTests(unittest.TestCase):
    def test_page_furniture_comparison_ignores_volatile_package_identity(self):
        before = {
            "section_count": 1,
            "even_and_odd_headers": False,
            "sections": [
                {
                    "section": 1,
                    "different_first_page": False,
                    "references": [
                        {
                            "kind": "header",
                            "type": "default",
                            "relationship": "rId8",
                            "part": "word/header1.xml",
                            "sha256": "AAA",
                            "visible_text": "Controlled title",
                            "fields": ["PAGE"],
                        }
                    ],
                }
            ],
        }
        after = {
            **before,
            "sections": [
                {
                    **before["sections"][0],
                    "references": [
                        {
                            **before["sections"][0]["references"][0],
                            "relationship": "rId243",
                            "part": "word/header_agdoc.xml",
                            "sha256": "BBB",
                        }
                    ],
                }
            ],
        }

        self.assertEqual(comparison_page_furniture(before), comparison_page_furniture(after))

    def test_reports_structural_visual_and_unexplained_change(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            before_docx = root / "before.docx"
            after_docx = root / "after.docx"
            before = Document()
            before.add_paragraph("Before")
            before.save(before_docx)
            after = Document()
            after.add_paragraph("Before")
            after.add_paragraph("Added")
            after.save(after_docx)

            before_page = root / "before.png"
            after_page = root / "after.png"
            Image.new("RGB", (40, 40), "white").save(before_page)
            changed = Image.new("RGB", (40, 40), "white")
            changed.putpixel((10, 10), (0, 0, 0))
            changed.save(after_page)

            result = compare_build_candidate(
                candidate_docx=after_docx,
                candidate_render={"page_images": [str(after_page)]},
                baseline_report={
                    "build_id": "before",
                    "run_directory": str(root),
                    "artifacts": {"docx": str(before_docx)},
                    "render": {"page_images": [str(before_page)]},
                },
                input_changes={"baseline_available": True, "changed": False, "items": []},
                output_directory=root / "diffs",
            )
            self.assertTrue(result["structure"]["changed"])
            self.assertTrue(result["visual"]["changed"])
            self.assertEqual(result["visual"]["changed_pages"], [1])
            self.assertTrue(result["unexplained_change"])
            self.assertTrue((root / "diffs" / "page-0001-difference.png").is_file())


if __name__ == "__main__":
    unittest.main()
