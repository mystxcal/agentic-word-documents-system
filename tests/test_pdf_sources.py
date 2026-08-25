import base64
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest
from unittest.mock import patch

from docx import Document

from agentic_docs.errors import PackageError
from agentic_docs.sources_pdf import compile_pdf_pages, expand_page_selection


class PdfSourceTests(unittest.TestCase):
    def test_expands_mixed_ordered_page_selectors(self):
        self.assertEqual(
            expand_page_selection([1, [3, 5], {"start": 9, "end": 10}]),
            [1, 3, 4, 5, 9, 10],
        )

    def test_rejects_invalid_page_ranges(self):
        with self.assertRaises(PackageError):
            expand_page_selection([[5, 3]])

    def test_compiles_declared_pages_into_native_word_fragment(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            style = root / "style.docx"
            Document().save(style)
            image = root / "page.png"
            image.write_bytes(
                base64.b64decode(
                    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
                )
            )
            source = root / "guide.pdf"
            source.write_bytes(b"pdf placeholder")
            output = root / "fragment.docx"
            rendered = {
                "pages": [
                    {"source_page": 2, "image": str(image)},
                    {"source_page": 3, "image": str(image)},
                ]
            }
            with patch("agentic_docs.sources_pdf.render_pdf_pages", return_value=rendered):
                result = compile_pdf_pages(
                    source,
                    output,
                    style_source=style,
                    options={"pages": [[2, 3]], "image_width_inches": 5.0},
                    title=None,
                    caption="Source page {page}",
                    alt_text="Guide page",
                    work_directory=root / "rendered",
                    available_width_inches=6.0,
                )
            compiled = Document(output)
            self.assertEqual(result["page_count"], 2)
            self.assertEqual(len(compiled.inline_shapes), 2)
            self.assertEqual(
                [paragraph.text for paragraph in compiled.paragraphs if paragraph.text],
                ["Source page 2", "Source page 3"],
            )


if __name__ == "__main__":
    unittest.main()
