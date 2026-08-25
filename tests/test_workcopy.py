import base64
from pathlib import Path
from types import SimpleNamespace
from tempfile import TemporaryDirectory
import unittest

from docx import Document

from agentic_docs.diagnostics import DiagnosticBag
from agentic_docs.build import _embed_component_baselines
from agentic_docs.errors import RevisionError
from agentic_docs.model import ComponentType, Ownership
from agentic_docs.word.ooxml import qn, wrap_elements
from agentic_docs.word.package import DOCUMENT_PART, DocxPackage, XNS, load_component_states, require_one_sdt
from agentic_docs.workcopy import _candidate_from_workcopy, adopt_workcopy, diff_workcopy


def tagged_docx(path: Path, tag: str, text: str) -> None:
    document = Document()
    paragraph = document.add_paragraph(text)
    wrap_elements(tag, [paragraph._p])
    document.save(path)


class WorkcopyTests(unittest.TestCase):
    def _resolved(self, root: Path, canonical: Path):
        declaration = SimpleNamespace(
            type=ComponentType.DOCUMENT,
            ownership=Ownership.WORD_FRAGMENT,
            source_tag="AGDOC.BODY.TEST",
        )
        component = SimpleNamespace(
            id="body",
            declaration=declaration,
            source_path=canonical,
        )
        manifest_path = root / "document.jsonc"
        manifest_path.write_text("{}", encoding="utf-8")
        return SimpleNamespace(
            components={"body": component},
            manifest_path=manifest_path,
            manifest=SimpleNamespace(id="example", outputs=SimpleNamespace(basename="Example")),
        )

    def test_diffs_and_adopts_tagged_component_without_build(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            canonical = root / "canonical.docx"
            workcopy = root / "workcopy.docx"
            tagged_docx(canonical, "AGDOC.BODY.TEST", "Original text")
            tagged_docx(workcopy, "AGDOC.COMPONENT.body", "Coworker text")
            resolved = self._resolved(root, canonical)

            difference = diff_workcopy(resolved, workcopy)
            self.assertFalse(difference["same"])
            self.assertEqual(difference["change_group_count"], 1)
            report = adopt_workcopy(
                resolved,
                DiagnosticBag(),
                workcopy,
                build=False,
            )
            self.assertTrue(Path(report["backup"]).is_file())
            package = DocxPackage(canonical)
            block = require_one_sdt(package.xml(DOCUMENT_PART), "AGDOC.BODY.TEST")
            text = "".join(block.xpath(".//w:t/text()", namespaces=XNS))
            self.assertEqual(text, "Coworker text")

    def test_embedded_baseline_detects_format_only_workcopy_change(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            canonical = root / "canonical.docx"
            compiled = root / "compiled.docx"
            workcopy = root / "workcopy.docx"
            tagged_docx(canonical, "AGDOC.BODY.TEST", "Same text")
            tagged_docx(compiled, "AGDOC.COMPONENT.body", "Same text")
            resolved = self._resolved(root, canonical)
            resolved.system_root = root
            resolved.manifest.sequence = [SimpleNamespace(items=["body"])]

            result = _embed_component_baselines(compiled, resolved, component_id=None)
            self.assertTrue(result["embedded"])
            self.assertIn("body", load_component_states(DocxPackage(compiled)))

            package = DocxPackage(compiled)
            block = require_one_sdt(package.xml(DOCUMENT_PART), "AGDOC.COMPONENT.body")
            run = block.xpath(".//w:r", namespaces=XNS)[0]
            run_properties = run.find("./w:rPr", namespaces=XNS)
            if run_properties is None:
                run_properties = run.makeelement(qn("w:rPr"))
                run.insert(0, run_properties)
            run_properties.append(run.makeelement(qn("w:b")))
            package.set_xml(DOCUMENT_PART, block.getroottree().getroot())
            package.write(workcopy)

            difference = diff_workcopy(resolved, workcopy)
            self.assertTrue(difference["content_same"])
            self.assertFalse(difference["formatting_same"])
            self.assertFalse(difference["same"])

    def test_adoption_candidate_prunes_media_from_unrelated_components(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            first_image = root / "first.png"
            second_image = root / "second.png"
            first_image.write_bytes(
                base64.b64decode(
                    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
                )
            )
            second_image.write_bytes(
                base64.b64decode(
                    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAIAAACQd1PeAAAADElEQVR42mP4z8AAAAMBAQDJ/pLvAAAAAElFTkSuQmCC"
                )
            )
            document = Document()
            first = document.add_paragraph()
            first.add_run().add_picture(str(first_image))
            second = document.add_paragraph()
            second.add_run().add_picture(str(second_image))
            wrap_elements("AGDOC.COMPONENT.body", [first._p])
            wrap_elements("AGDOC.COMPONENT.other", [second._p])
            workcopy = root / "workcopy.docx"
            candidate = root / "candidate.docx"
            document.save(workcopy)

            _candidate_from_workcopy(
                workcopy,
                candidate,
                canonical_tag="AGDOC.BODY.TEST",
                workcopy_tag="AGDOC.COMPONENT.body",
            )
            package = DocxPackage(candidate)
            media = [name for name in package.parts if name.startswith("word/media/")]
            self.assertEqual(len(media), 1)

    def test_adoption_refuses_independent_canonical_and_workcopy_content_changes(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            canonical = root / "canonical.docx"
            compiled = root / "compiled.docx"
            workcopy = root / "workcopy.docx"
            tagged_docx(canonical, "AGDOC.BODY.TEST", "Baseline text")
            tagged_docx(compiled, "AGDOC.COMPONENT.body", "Baseline text")
            resolved = self._resolved(root, canonical)
            resolved.system_root = root
            resolved.manifest.sequence = [SimpleNamespace(items=["body"])]
            _embed_component_baselines(compiled, resolved, component_id=None)

            tagged_docx(canonical, "AGDOC.BODY.TEST", "New canonical text")
            package = DocxPackage(compiled)
            block = require_one_sdt(package.xml(DOCUMENT_PART), "AGDOC.COMPONENT.body")
            text_node = block.xpath(".//w:t", namespaces=XNS)[0]
            text_node.text = "Returned coworker text"
            package.set_xml(DOCUMENT_PART, block.getroottree().getroot())
            package.write(workcopy)

            difference = diff_workcopy(resolved, workcopy)
            self.assertTrue(difference["content_conflict"])
            with self.assertRaisesRegex(RevisionError, "both changed"):
                adopt_workcopy(resolved, DiagnosticBag(), workcopy, build=False)


if __name__ == "__main__":
    unittest.main()
