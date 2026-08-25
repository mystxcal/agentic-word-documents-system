from __future__ import annotations

import io
import tempfile
import unittest
from pathlib import Path

from docx import Document
from PIL import Image

from agentic_docs.word.package import DocxPackage, XNS, normalize_imported_media, set_do_not_compress_pictures


class ImportedMediaNormalizationTests(unittest.TestCase):
    def test_indexed_png_is_expanded_deterministically(self):
        image = Image.new("P", (8, 8))
        image.putpalette([0, 0, 0, 255, 255, 255] + [0, 0, 0] * 254)
        image.putdata([0, 1] * 32)
        original = io.BytesIO()
        image.save(original, format="PNG")

        first, report = normalize_imported_media("word/media/logo.png", original.getvalue())
        second, second_report = normalize_imported_media("word/media/logo.png", original.getvalue())

        self.assertEqual(first, second)
        self.assertEqual(report, second_report)
        self.assertEqual("RGB", Image.open(io.BytesIO(first)).mode)
        self.assertEqual("indexed-png-expanded-for-stable-word-rendering", report["reason"])

    def test_expanded_png_passes_through_byte_for_byte(self):
        image = Image.new("RGBA", (8, 8), (20, 40, 60, 180))
        original = io.BytesIO()
        image.save(original, format="PNG")

        result, report = normalize_imported_media("word/media/logo.png", original.getvalue())

        self.assertEqual(original.getvalue(), result)
        self.assertIsNone(report)

    def test_package_records_normalization_once_per_source_part(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            image_path = root / "indexed.png"
            image = Image.new("P", (16, 16))
            image.putpalette([10, 20, 30] * 256)
            image.save(image_path, format="PNG")

            donor_path = root / "donor.docx"
            donor = Document()
            donor.add_picture(str(image_path))
            donor.save(donor_path)

            target_path = root / "target.docx"
            Document().save(target_path)

            source = DocxPackage(donor_path)
            target = DocxPackage(target_path)
            source_part = next(name for name in source.parts if name.startswith("word/media/"))

            copied = target.copy_part_from(source, source_part)
            copied_again = target.copy_part_from(source, source_part)

            self.assertEqual(copied, copied_again)
            self.assertEqual(1, len(target.media_normalizations))
            self.assertEqual("RGB", Image.open(io.BytesIO(target.parts[copied])).mode)

    def test_forced_header_clones_receive_disjoint_drawing_ids(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            image_path = root / "logo.png"
            Image.new("RGB", (24, 12), (20, 80, 140)).save(image_path, format="PNG")

            donor_path = root / "donor.docx"
            donor = Document()
            donor.sections[0].header.paragraphs[0].add_run().add_picture(str(image_path))
            donor.save(donor_path)

            target_path = root / "target.docx"
            Document().save(target_path)
            source = DocxPackage(donor_path)
            target = DocxPackage(target_path)
            source_part = next(name for name in source.parts if name.startswith("word/header"))

            first = target.copy_part_from(source, source_part, force_unique=True)
            second = target.copy_part_from(source, source_part, force_unique=True)
            first_ids = {
                int(node.get("id"))
                for node in target.xml(first).xpath(".//wp:docPr|.//pic:cNvPr", namespaces=XNS)
            }
            second_ids = {
                int(node.get("id"))
                for node in target.xml(second).xpath(".//wp:docPr|.//pic:cNvPr", namespaces=XNS)
            }

            self.assertNotEqual(first, second)
            self.assertTrue(first_ids)
            self.assertTrue(second_ids)
            self.assertTrue(first_ids.isdisjoint(second_ids))

    def test_word_picture_compression_policy_is_explicit_and_idempotent(self):
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "document.docx"
            Document().save(path)
            package = DocxPackage(path)

            first = set_do_not_compress_pictures(package, True)
            second = set_do_not_compress_pictures(package, True)
            nodes = package.xml("word/settings.xml").xpath(
                "./w:doNotCompressPictures[@w:val='true']",
                namespaces=XNS,
            )

            self.assertTrue(first["changed"])
            self.assertFalse(second["changed"])
            self.assertEqual(1, len(nodes))


if __name__ == "__main__":
    unittest.main()
