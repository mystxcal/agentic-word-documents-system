from pathlib import Path
from tempfile import TemporaryDirectory
import unittest
import zipfile

from docx import Document

from agentic_docs.docx_text import TextReplacement, replace_docx_text, replace_plain_text


class DocxTextReplacementTests(unittest.TestCase):
    def test_replaces_across_runs_without_rebuilding_package(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.docx"
            output = root / "output.docx"
            document = Document()
            paragraph = document.add_paragraph()
            first = paragraph.add_run("Copper fi")
            first.bold = True
            second = paragraph.add_run("bre links and fibre.")
            second.italic = True
            document.save(source)

            summary = replace_docx_text(
                source,
                output,
                [TextReplacement("fibre", "fiber", whole_word=True)],
            )

            self.assertEqual(summary.total_replacements, 2)
            self.assertEqual(summary.parts_changed, ["word/document.xml"])
            reopened = Document(output)
            self.assertEqual(reopened.paragraphs[0].text, "Copper fiber links and fiber.")
            self.assertTrue(reopened.paragraphs[0].runs[0].bold)
            self.assertTrue(reopened.paragraphs[0].runs[1].italic)
            with zipfile.ZipFile(source) as before, zipfile.ZipFile(output) as after:
                self.assertEqual(set(before.namelist()), set(after.namelist()))
                changed = [
                    name
                    for name in before.namelist()
                    if before.read(name) != after.read(name)
                ]
                self.assertEqual(changed, ["word/document.xml"])

    def test_case_insensitive_replacement_is_explicit(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.docx"
            output = root / "output.docx"
            document = Document()
            document.add_paragraph("Fibre fibre FIBRE")
            document.save(source)
            summary = replace_docx_text(
                source,
                output,
                [TextReplacement("fibre", "fiber", case_sensitive=False, whole_word=True)],
            )
            self.assertEqual(summary.total_replacements, 3)
            self.assertEqual(Document(output).paragraphs[0].text, "fiber fiber fiber")

    def test_replaces_canonical_markdown_without_destroying_utf8_text(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.md"
            output = root / "output.md"
            source.write_text("# Field note\n\nUse fibre.\n\nملاحظة", encoding="utf-8")

            summary = replace_plain_text(
                source,
                output,
                [TextReplacement("fibre", "fiber", whole_word=True)],
            )

            self.assertEqual(summary.total_replacements, 1)
            self.assertEqual(output.read_text(encoding="utf-8"), "# Field note\n\nUse fiber.\n\nملاحظة")


if __name__ == "__main__":
    unittest.main()
