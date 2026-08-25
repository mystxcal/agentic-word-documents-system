from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document
from lxml import etree

from agentic_docs.diagnostics import DiagnosticBag
from agentic_docs.word.fragments import COMMENTS_REL, component_wrapper, import_word_fragment, replace_nested_slot
from agentic_docs.word.ooxml import CT_NS, W_NS, qn, wrap_elements
from agentic_docs.word.package import (
    CONTENT_TYPES,
    DOCUMENT_PART,
    DocxPackage,
    XNS,
    resolve_target,
    validate_docx,
)


def _commented_fragment(path: Path) -> None:
    plain = path.with_name("plain.docx")
    document = Document()
    paragraph = document.add_paragraph("Reviewed text")
    wrap_elements("AGDOC.BODY.TEST", [paragraph._p])
    document.save(plain)

    package = DocxPackage(plain)
    root = package.xml(DOCUMENT_PART)
    paragraph_node = root.xpath(".//w:sdt[w:sdtPr/w:tag[@w:val='AGDOC.BODY.TEST']]//w:p", namespaces=XNS)[0]
    run = paragraph_node.xpath("./w:r", namespaces=XNS)[0]
    start = etree.Element(qn("w:commentRangeStart"))
    start.set(qn("w:id"), "0")
    end = etree.Element(qn("w:commentRangeEnd"))
    end.set(qn("w:id"), "0")
    reference_run = etree.Element(qn("w:r"))
    reference = etree.SubElement(reference_run, qn("w:commentReference"))
    reference.set(qn("w:id"), "0")
    run_index = paragraph_node.index(run)
    paragraph_node.insert(run_index, start)
    paragraph_node.insert(run_index + 2, end)
    paragraph_node.insert(run_index + 3, reference_run)
    package.set_xml(DOCUMENT_PART, root)

    comments = etree.Element(qn("w:comments"), nsmap={"w": W_NS})
    comment = etree.SubElement(comments, qn("w:comment"))
    comment.set(qn("w:id"), "0")
    comment.set(qn("w:author"), "Reviewer")
    comment.set(qn("w:initials"), "RV")
    comment_paragraph = etree.SubElement(comment, qn("w:p"))
    comment_run = etree.SubElement(comment_paragraph, qn("w:r"))
    comment_text = etree.SubElement(comment_run, qn("w:t"))
    comment_text.text = "Please confirm this sentence."
    package.set_xml("word/comments.xml", comments)
    package.add_relationship(DOCUMENT_PART, rel_type=COMMENTS_REL, target="comments.xml")

    content_types = package.xml(CONTENT_TYPES)
    etree.SubElement(
        content_types,
        f"{{{CT_NS}}}Override",
        PartName="/word/comments.xml",
        ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
    )
    package.set_xml(CONTENT_TYPES, content_types)
    package.write(path)


class WordFragmentTests(unittest.TestCase):
    def test_replaces_explicit_nested_word_slot_without_touching_surrounding_text(self):
        document = Document()
        before = document.add_paragraph("Before")
        placeholder = document.add_paragraph("Placeholder")
        after = document.add_paragraph("After")
        slot = wrap_elements("AGDOC.SLOT.TABLE", [placeholder._p])
        parent = component_wrapper("parent", [before._p, slot, after._p])
        child_paragraph = document.add_paragraph("Inserted child")._p
        replace_nested_slot(parent, "AGDOC.SLOT.TABLE", [component_wrapper("child", [child_paragraph])])
        text = parent.xpath(".//w:t/text()", namespaces=XNS)
        self.assertEqual(text, ["Before", "Inserted child", "After"])

    def test_import_preserves_referenced_word_comment(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.docx"
            target_path = root / "target.docx"
            output = root / "output.docx"
            _commented_fragment(source)
            Document().save(target_path)

            target = DocxPackage(target_path)
            wrapper = import_word_fragment(
                target,
                source,
                component_id="body",
                source_tag="AGDOC.BODY.TEST",
                allow_untagged=False,
                preserve_sections=False,
                diagnostics=DiagnosticBag(),
            )
            document_root = target.xml(DOCUMENT_PART)
            body = document_root.find(qn("w:body"))
            section = body.find(qn("w:sectPr"))
            body.insert(body.index(section), wrapper)
            target.set_xml(DOCUMENT_PART, document_root)
            target.write(output)

            self.assertTrue(validate_docx(output)["valid"])
            result = DocxPackage(output)
            _, relationships = result.relationship_root(DOCUMENT_PART)
            comment_relationships = relationships.xpath(
                "./pr:Relationship[@Type=$value]", namespaces=XNS, value=COMMENTS_REL
            )
            self.assertEqual(len(comment_relationships), 1)
            comment_part = resolve_target(DOCUMENT_PART, comment_relationships[0].get("Target"))
            comment_root = result.xml(comment_part)
            self.assertEqual(
                "".join(comment_root.xpath(".//w:comment[@w:id='0']//w:t/text()", namespaces=XNS)),
                "Please confirm this sentence.",
            )
            self.assertEqual(
                wrapper.xpath("string(.//w:commentReference/@w:id)", namespaces=XNS),
                "0",
            )


if __name__ == "__main__":
    unittest.main()
