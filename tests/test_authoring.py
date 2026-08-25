import json
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest
from unittest.mock import patch

from docx import Document
from openpyxl import Workbook
from openpyxl.worksheet.table import Table
from PIL import Image

from agentic_docs.authoring import (
    accept_diagram_rendition,
    add_document_component,
    create_document,
    create_project,
    parse_pages,
)
from agentic_docs.diagnostics import DiagnosticBag
from agentic_docs.errors import DocumentSystemError, ResolutionError
from agentic_docs.jsonc import load_jsonc
from agentic_docs.guidance import document_workspace, edit_target
from agentic_docs.model import DocumentManifest
from agentic_docs.resolver import file_hash, resolve_document
from agentic_docs.word.components import compile_component_wrapper
from agentic_docs.word.ooxml import wrap_elements
from agentic_docs.word.package import DocxPackage, validate_docx


class AuthoringTests(unittest.TestCase):
    @staticmethod
    def _png(path: Path) -> None:
        Image.new("RGB", (320, 180), "white").save(path, format="PNG")

    def _system(self, root: Path):
        kit = root / "kits" / "studio"
        donor = kit / "donors" / "styles.docx"
        donor.parent.mkdir(parents=True)
        Document().save(donor)
        (kit / "kit.jsonc").write_text(
            json.dumps(
                {
                    "schema": "agentic-kit/v2",
                    "id": "studio",
                    "components": {"studio": "donors/styles.docx"},
                    "semantic_styles": {"heading_1": "Heading 1", "body": "Normal"},
                    "table_styles": {
                        "technical": {
                            "font_name": "Calibri",
                            "font_size_pt": 8.5,
                            "header_fill": "1F3763",
                            "header_text": "FFFFFF",
                            "group_fill": "4472C4",
                            "group_text": "FFFFFF",
                            "body_fill": "FFFFFF",
                            "alternate_fill": "F2F6FC",
                            "border_color": "B4C6E7",
                            "text_color": "1F1F1F",
                            "cell_margin_twips": 90,
                        }
                    },
                }
            ),
            encoding="utf-8",
        )
        profile = root / "profiles" / "plain"
        profile.mkdir(parents=True)
        shell = Document()
        paragraph = shell.add_paragraph("Body")
        wrap_elements("AGDOC.BODY.PLAIN", [paragraph._p])
        shell.save(profile / "shell.docx")
        (profile / "profile.jsonc").write_text(
            json.dumps(
                {
                    "schema": "agentic-profile/v2",
                    "id": "plain",
                    "shell": "shell.docx",
                    "body_slot": "AGDOC.BODY.PLAIN",
                    "field_bindings": {},
                    "release_gates": [],
                }
            ),
            encoding="utf-8",
        )
        create_project(root, "demo", name="Demo", description="Neutral test collection")
        result = create_document(
            root,
            "demo-doc",
            project_id="demo",
            title="Demo Document",
            document_type="Engineering Document",
            revision="R00",
        )
        return resolve_document(Path(result["manifest"]))

    def test_creates_buildable_markdown_first_document_model(self):
        with TemporaryDirectory() as directory:
            resolved = self._system(Path(directory))
            self.assertTrue(resolved.manifest_path.is_file())
            self.assertEqual(resolved.components["body"].source_path.read_text(encoding="utf-8"), "# Demo Document\n")
            self.assertIsInstance(DocumentManifest.model_validate(load_jsonc(resolved.manifest_path)), DocumentManifest)

    def test_workspace_selects_primary_canonical_source_and_gives_exact_commands(self):
        with TemporaryDirectory() as directory:
            resolved = self._system(Path(directory))
            workspace = document_workspace(
                resolved,
                {
                    "build_id": None,
                    "ready": False,
                    "sources_current": False,
                    "current_integrity": False,
                    "build_quality_passed": False,
                    "verification_mode": None,
                },
            )

            self.assertEqual("body", workspace["primary_edit"]["id"])
            self.assertEqual("markdown-plus", workspace["primary_edit"]["source_kind"])
            self.assertEqual("not-built", workspace["next_action"]["state"])
            self.assertIn("build demo-doc --quick", workspace["next_action"]["command"])
            self.assertTrue(any(item["role"] == "styles" for item in workspace["presentation"]))

    def test_edit_locates_primary_or_shared_presentation_without_opening(self):
        with TemporaryDirectory() as directory:
            resolved = self._system(Path(directory))

            primary = edit_target(resolved, open_file=False)
            styles = edit_target(resolved, presentation="styles", open_file=False)

            self.assertEqual("body", primary["selection"])
            self.assertEqual(resolved.components["body"].source_path, Path(primary["path"]))
            self.assertEqual("shared-kit", styles["scope"])
            self.assertIsNotNone(styles["shared_warning"])

    def test_creates_transactional_word_owned_document_without_markdown_conversion(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            self._system(root)
            source = root / "received-handbook.docx"
            document = Document()
            document.add_heading("Received Handbook", level=1)
            document.add_paragraph("Word-owned editorial content")
            document.save(source)
            result = create_document(
                root,
                "word-owned",
                project_id="demo",
                title="Word Owned Handbook",
                document_type="Handbook",
                revision="R00",
                word_source=source,
                allow_untagged=True,
                preserve_sections=True,
                use_source_styles=True,
                preserve_source_layout=True,
            )
            resolved = resolve_document(Path(result["manifest"]))
            body = resolved.components["body"]
            self.assertEqual(body.declaration.ownership.value, "word_fragment")
            self.assertTrue(body.declaration.options.preserve_sections)
            self.assertTrue(body.declaration.options.whole_document)
            self.assertEqual(resolved.presentation_paths["styles"], body.source_path)
            self.assertEqual(resolved.manifest.presentation.page_regions["main"].layout_mode, "preserve")
            self.assertIsNone(resolved.manifest.presentation.page_regions["main"].numbering)
            self.assertTrue(result["source_layout_preserved"])
            self.assertEqual(body.source_hash, file_hash(source))
            self.assertFalse((Path(result["root"]) / "content" / "body.md").exists())

    def test_adds_excel_component_and_preserves_one_authorized_placement(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            resolved = self._system(root)
            body = resolved.components["body"].source_path
            body.write_text(body.read_text(encoding="utf-8") + "\n:::insert reading-list\n", encoding="utf-8")
            workbook_path = root / "received-reading-list.xlsx"
            workbook = Workbook()
            worksheet = workbook.active
            worksheet.title = "Reading list"
            worksheet.append(["Title", "Pages"])
            worksheet.append(["Field Notes", 12])
            worksheet.add_table(Table(displayName="ReadingListData", ref="A1:B2"))
            workbook.save(workbook_path)
            result = add_document_component(
                resolve_document(resolved.manifest_path),
                kind="table",
                component_id="reading-list-data",
                source=workbook_path,
                parent_component="body",
                slot_name="reading-list",
                excel_table="ReadingListData",
            )
            updated = resolve_document(resolved.manifest_path)
            self.assertEqual(updated.manifest.components["body"].slots["reading-list"], ["reading-list-data"])
            self.assertTrue(updated.components["reading-list-data"].source_path.is_file())
            self.assertTrue(Path(result["recovery"]).is_dir())

    def test_missing_markdown_marker_refuses_partial_manifest_change(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            resolved = self._system(root)
            original = resolved.manifest_path.read_bytes()
            image = root / "diagram.png"
            image.write_bytes(
                bytes.fromhex(
                    "89504E470D0A1A0A0000000D49484452000000010000000108060000001F15C489"
                    "0000000D49444154789C6360000000020001E221BC330000000049454E44AE426082"
                )
            )
            with self.assertRaisesRegex(DocumentSystemError, ":::insert diagram"):
                add_document_component(
                    resolved,
                    kind="figure",
                    component_id="diagram",
                    source=image,
                    parent_component="body",
                    alt_text="Process diagram",
                )
            self.assertEqual(resolved.manifest_path.read_bytes(), original)

    def test_adds_native_diagram_and_refuses_a_stale_rendition(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            resolved = self._system(root)
            body = resolved.components["body"].source_path
            body.write_text(body.read_text(encoding="utf-8") + "\n:::insert journey-map\n", encoding="utf-8")
            native = root / "journey-map.drawio"
            native.write_text("<mxfile revision='one'/>", encoding="utf-8")
            rendition = root / "journey-map.png"
            self._png(rendition)

            result = add_document_component(
                resolved,
                kind="diagram",
                component_id="journey-map",
                source=native,
                rendition=rendition,
                parent_component="body",
                alt_text="Journey map",
            )
            updated = resolve_document(resolved.manifest_path)
            diagram = updated.components["journey-map"]
            self.assertEqual(diagram.declaration.type.value, "diagram")
            self.assertEqual(diagram.source_path.suffix, ".drawio")
            self.assertEqual(diagram.related_paths["rendition"].suffix, ".png")
            self.assertEqual(
                diagram.source_hash,
                diagram.declaration.options.rendition_of_sha256,
            )
            self.assertIn("rendition", result["related_sources"])

            package = DocxPackage(updated.shell_path)
            wrapper = compile_component_wrapper(
                package,
                updated,
                diagram,
                build_work=root / "compile-work",
                diagnostics=DiagnosticBag(),
            )
            self.assertIsNotNone(wrapper)
            compiled = root / "diagram-publication.docx"
            package.write(compiled)
            validate_docx(compiled)

            diagram.source_path.write_text("<mxfile revision='two'/>", encoding="utf-8")
            with self.assertRaisesRegex(ResolutionError, "accept-rendition"):
                resolve_document(resolved.manifest_path)

            refreshed = root / "journey-map-reviewed.png"
            self._png(refreshed)
            stale = resolve_document(resolved.manifest_path, allow_stale_diagrams=True)
            receipt = accept_diagram_rendition(stale, component_id="journey-map", rendition=refreshed)
            accepted = resolve_document(resolved.manifest_path)
            self.assertEqual(
                accepted.components["journey-map"].source_hash,
                accepted.components["journey-map"].declaration.options.rendition_of_sha256,
            )
            self.assertEqual(receipt["operation"], "accept-diagram-rendition")

    def test_failed_diagram_commit_removes_both_new_canonical_files(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            resolved = self._system(root)
            body = resolved.components["body"].source_path
            body.write_text(body.read_text(encoding="utf-8") + "\n:::insert journey-map\n", encoding="utf-8")
            native = root / "journey-map.vsdx"
            native.write_bytes(b"native")
            rendition = root / "journey-map.png"
            self._png(rendition)
            with patch("agentic_docs.resolver.resolve_document", side_effect=DocumentSystemError("synthetic failure")):
                with self.assertRaisesRegex(DocumentSystemError, "synthetic failure"):
                    add_document_component(
                        resolved,
                        kind="diagram",
                        component_id="journey-map",
                        source=native,
                        rendition=rendition,
                        parent_component="body",
                        alt_text="Journey map",
                    )
            sources = resolved.manifest_path.parent / "content" / "sources"
            self.assertFalse((sources / "journey-map.vsdx").exists())
            self.assertFalse((sources / "journey-map.png").exists())

    def test_parses_human_pdf_page_selection(self):
        self.assertEqual(parse_pages("1, 3-5, 8"), [1, [3, 5], 8])

    def test_failed_document_creation_removes_the_entire_starter(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            resolved = self._system(root)
            target = root / "projects" / "demo" / "documents" / "broken-doc"
            with patch("agentic_docs.resolver.resolve_document", side_effect=DocumentSystemError("synthetic failure")):
                with self.assertRaisesRegex(DocumentSystemError, "synthetic failure"):
                    create_document(
                        root,
                        "broken-doc",
                        project_id="demo",
                        title="Broken Document",
                        document_type="Engineering Document",
                        revision="R00",
                    )
            self.assertFalse(target.exists())
            self.assertEqual(list(target.parent.glob(".broken-doc.creating-*")), [])

    def test_failed_component_commit_restores_manifest_marker_and_new_source(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            resolved = self._system(root)
            body = resolved.components["body"].source_path
            body.write_text(body.read_text(encoding="utf-8") + "\n:::insert diagram\n", encoding="utf-8")
            original_manifest = resolved.manifest_path.read_bytes()
            original_body = body.read_bytes()
            image = root / "diagram.png"
            image.write_bytes(
                bytes.fromhex(
                    "89504E470D0A1A0A0000000D49484452000000010000000108060000001F15C489"
                    "0000000D49444154789C6360000000020001E221BC330000000049454E44AE426082"
                )
            )
            with patch("agentic_docs.resolver.resolve_document", side_effect=DocumentSystemError("synthetic failure")):
                with self.assertRaisesRegex(DocumentSystemError, "synthetic failure"):
                    add_document_component(
                        resolved,
                        kind="figure",
                        component_id="diagram",
                        source=image,
                        parent_component="body",
                        alt_text="Process diagram",
                    )
            self.assertEqual(resolved.manifest_path.read_bytes(), original_manifest)
            self.assertEqual(body.read_bytes(), original_body)
            self.assertFalse((resolved.manifest_path.parent / "content" / "sources" / "diagram.png").exists())
            rollbacks = list((resolved.manifest_path.parent / ".history" / "authoring").glob("*/rolled-back.json"))
            self.assertEqual(len(rollbacks), 1)


if __name__ == "__main__":
    unittest.main()
