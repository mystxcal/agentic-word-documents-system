"""Generate the binary Word, Excel, and image assets shipped with the examples.

The human-authored manifests and Markdown stay outside this script. This file
owns only reproducible binary donors and sample data assets.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.worksheet.table import Table, TableStyleInfo
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from agentic_docs.word.ooxml import append_complex_field, wrap_elements  # noqa: E402


THEMES = {
    "studio": {
        "ink": "18242B",
        "primary": "17324D",
        "accent": "1E7A78",
        "warm": "D8A84E",
        "pale": "EDF4F3",
        "paper": "FFFFFF",
        "muted": "62737B",
        "callout": "E4F0EF",
        "font": "Aptos",
        "code_font": "Cascadia Mono",
        "callout_style": "Studio Callout",
        "code_style": "Studio Code",
        "header_label": "FIELD REPORT",
    },
    "storyboard": {
        "ink": "2A2430",
        "primary": "3D2B56",
        "accent": "D75D45",
        "warm": "EFCB68",
        "pale": "F7EFE8",
        "paper": "FFFDF9",
        "muted": "746777",
        "callout": "F3E5D8",
        "font": "Aptos",
        "code_font": "Cascadia Mono",
        "callout_style": "Storyboard Callout",
        "code_style": "Storyboard Cue",
        "header_label": "PRACTICE GUIDE",
    },
}


def _rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def _set_cell_fill(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top=100, start=120, bottom=100, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table, color: str, size: int = 6, *, top=True, bottom=True, inside=True) -> None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    names = ["left", "right"]
    if top:
        names.append("top")
    if bottom:
        names.append("bottom")
    if inside:
        names.extend(["insideH", "insideV"])
    for name in names:
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:color"), color)


def _set_paragraph_bottom_border(paragraph, color: str, size: int = 10) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _configure_page(document: Document, *, margins=(1.7, 1.8, 1.7, 1.8)) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(margins[0])
    section.right_margin = Cm(margins[1])
    section.bottom_margin = Cm(margins[2])
    section.left_margin = Cm(margins[3])
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)


def _set_font(style, name: str, size: float, color: str, *, bold=None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = _rgb(color)
    if bold is not None:
        style.font.bold = bold
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _set_document_properties(document: Document, *, title: str, subject: str) -> None:
    document.core_properties.title = title
    document.core_properties.subject = subject
    document.core_properties.author = "Agentic Word Documents contributors"
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = "Generated from tools/build_template_assets.py"


def _paragraph_shading(style, fill: str, accent: str) -> None:
    properties = style.element.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    properties.append(borders)


def create_style_donor(path: Path, theme_name: str) -> None:
    theme = THEMES[theme_name]
    document = Document()
    _configure_page(document)
    styles = document.styles

    normal = styles["Normal"]
    _set_font(normal, theme["font"], 10.5, theme["ink"])
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.13

    heading_specs = {
        "Heading 1": (23, theme["primary"], 24, 10),
        "Heading 2": (15.5, theme["accent"], 18, 6),
        "Heading 3": (11.5, theme["primary"], 13, 4),
        "Heading 4": (10.5, theme["muted"], 10, 3),
        "Heading 5": (9.5, theme["muted"], 8, 2),
        "Heading 6": (9.0, theme["muted"], 8, 2),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        _set_font(style, theme["font"], size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    styles["Heading 1"].paragraph_format.page_break_before = False

    title = styles["Title"]
    _set_font(title, theme["font"], 34 if theme_name == "studio" else 37, theme["primary"], bold=True)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.keep_with_next = True
    subtitle = styles["Subtitle"]
    _set_font(subtitle, theme["font"], 13, theme["muted"])
    subtitle.paragraph_format.space_after = Pt(12)

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        _set_font(style, theme["font"], 10.5, theme["ink"])
        style.paragraph_format.space_after = Pt(3)

    caption = styles["Caption"]
    _set_font(caption, theme["font"], 8.5, theme["muted"])
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)

    callout = styles.add_style(theme["callout_style"], WD_STYLE_TYPE.PARAGRAPH)
    callout.base_style = normal
    _set_font(callout, theme["font"], 10, theme["ink"])
    callout.paragraph_format.left_indent = Cm(0.35)
    callout.paragraph_format.right_indent = Cm(0.2)
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(8)
    _paragraph_shading(callout, theme["callout"], theme["accent"])

    code = styles.add_style(theme["code_style"], WD_STYLE_TYPE.PARAGRAPH)
    code.base_style = normal
    _set_font(code, theme["code_font"], 9.0, theme["primary"])
    code.paragraph_format.left_indent = Cm(0.45)
    code.paragraph_format.right_indent = Cm(0.2)
    code.paragraph_format.space_before = Pt(5)
    code.paragraph_format.space_after = Pt(7)
    _paragraph_shading(code, theme["pale"], theme["warm"])

    _set_document_properties(
        document,
        title=f"{theme_name.title()} style donor",
        subject="Reusable Word-native style kit",
    )
    _clear_body(document)
    document.add_paragraph("Style donor — source asset, not a published page.")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _tag_paragraph(paragraph, tag: str, alias: str) -> None:
    wrap_elements(tag, [paragraph._p], alias=alias)


def create_cover(path: Path, style_source: Path, theme_name: str) -> None:
    theme = THEMES[theme_name]
    document = Document(str(style_source))
    _set_document_properties(
        document,
        title=f"{theme_name.title()} sample cover",
        subject="Editable document-level cover fragment",
    )
    _clear_body(document)
    _configure_page(document, margins=(1.1, 1.25, 1.1, 1.25))

    stripe = document.add_table(rows=1, cols=1)
    stripe.alignment = WD_TABLE_ALIGNMENT.CENTER
    stripe.autofit = False
    stripe.columns[0].width = Cm(18.2)
    stripe.rows[0].height = Cm(0.35)
    stripe.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    _set_cell_fill(stripe.cell(0, 0), theme["accent"])
    _set_table_borders(stripe, theme["accent"], 0, top=False, bottom=False, inside=False)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(62 if theme_name == "studio" else 48)

    kicker = document.add_paragraph("DOCUMENT TYPE")
    kicker.style = "Subtitle"
    kicker.paragraph_format.space_after = Pt(10)
    kicker.runs[0].font.bold = True
    kicker.runs[0].font.color.rgb = _rgb(theme["accent"])
    kicker.runs[0].font.size = Pt(9.5)
    _tag_paragraph(kicker, "AGDOC.FIELD.COVER.KICKER", "Cover kicker")

    title = document.add_paragraph("DOCUMENT TITLE")
    title.style = "Title"
    title.paragraph_format.space_after = Pt(14)
    _tag_paragraph(title, "AGDOC.FIELD.COVER.TITLE", "Cover title")

    rule = document.add_table(rows=1, cols=2)
    rule.alignment = WD_TABLE_ALIGNMENT.LEFT
    rule.autofit = False
    rule.columns[0].width = Cm(3.2)
    rule.columns[1].width = Cm(14.8)
    rule.rows[0].height = Cm(0.10)
    rule.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    _set_cell_fill(rule.cell(0, 0), theme["warm"])
    _set_cell_fill(rule.cell(0, 1), theme["pale"])
    _set_table_borders(rule, theme["paper"], 0, top=False, bottom=False, inside=False)

    subtitle = document.add_paragraph("A short description of this document.")
    subtitle.style = "Subtitle"
    subtitle.paragraph_format.space_before = Pt(16)
    subtitle.paragraph_format.space_after = Pt(40)
    _tag_paragraph(subtitle, "AGDOC.FIELD.COVER.SUBTITLE", "Cover subtitle")

    meta = document.add_table(rows=2, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    meta.columns[0].width = Cm(8.8)
    meta.columns[1].width = Cm(8.8)
    labels = (("AUTHOR", "AUTHOR NAME"), ("ISSUED", "DATE"))
    tags = ("AGDOC.FIELD.COVER.AUTHOR", "AGDOC.FIELD.COVER.DATE")
    for row_index, ((label, value), tag) in enumerate(zip(labels, tags)):
        left = meta.cell(row_index, 0)
        right = meta.cell(row_index, 1)
        for cell in (left, right):
            _set_cell_margins(cell, top=95, bottom=95, start=80, end=80)
        label_p = left.paragraphs[0]
        label_p.text = label
        label_p.runs[0].font.name = theme["font"]
        label_p.runs[0].font.size = Pt(8)
        label_p.runs[0].font.bold = True
        label_p.runs[0].font.color.rgb = _rgb(theme["muted"])
        value_p = right.paragraphs[0]
        value_p.text = value
        value_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        value_p.runs[0].font.name = theme["font"]
        value_p.runs[0].font.size = Pt(9.5)
        value_p.runs[0].font.bold = True
        value_p.runs[0].font.color.rgb = _rgb(theme["primary"])
        _tag_paragraph(value_p, tag, label.title())
    _set_table_borders(meta, theme["warm"], 5, top=True, bottom=True, inside=True)

    bottom_space = document.add_paragraph()
    bottom_space.paragraph_format.space_after = Pt(64 if theme_name == "studio" else 52)
    note = document.add_paragraph("EDITABLE TEMPLATE  ·  CANONICAL SOURCES  ·  REPRODUCIBLE OUTPUT")
    note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = note.runs[0]
    run.font.name = theme["font"]
    run.font.size = Pt(7.5)
    run.font.bold = True
    run.font.color.rgb = _rgb(theme["muted"])

    body_elements = [child for child in list(document._element.body) if child.tag != qn("w:sectPr")]
    wrap_elements("AGDOC.COVER", body_elements, alias="Document cover")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def create_shell(path: Path, style_source: Path, body_tag: str) -> None:
    document = Document(str(style_source))
    _set_document_properties(
        document,
        title=f"{body_tag.rsplit('.', 1)[-1].title()} profile shell",
        subject="Reusable document structure",
    )
    _clear_body(document)
    _configure_page(document)

    cover = document.add_paragraph("COVER")
    wrap_elements("AGDOC.COVER", [cover._p], alias="Replaceable cover")

    toc_heading = document.add_paragraph("Contents")
    toc_heading.style = "TOC Heading"
    toc_field = document.add_paragraph()
    append_complex_field(toc_field, 'TOC \\o "1-2" \\h \\z \\u', "Update fields to populate the table of contents")
    toc_field.paragraph_format.space_after = Pt(18)
    wrap_elements(
        "AGDOC.LAYOUT.MAIN_START",
        [toc_heading._p, toc_field._p],
        alias="Main document start and table of contents",
    )

    placeholder = document.add_paragraph("CANONICAL BODY")
    wrap_elements(body_tag, [placeholder._p], alias="Canonical document body")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def create_plain_shell(path: Path, style_source: Path) -> None:
    document = Document(str(style_source))
    _set_document_properties(document, title="Plain profile shell", subject="Reusable document structure")
    _clear_body(document)
    _configure_page(document)
    placeholder = document.add_paragraph("CANONICAL BODY")
    wrap_elements("AGDOC.BODY.PLAIN", [placeholder._p], alias="Canonical document body")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def create_header(path: Path, style_source: Path, theme_name: str) -> None:
    theme = THEMES[theme_name]
    document = Document(str(style_source))
    _set_document_properties(
        document,
        title=f"{theme_name.title()} running header",
        subject="Reusable Word-native header donor",
    )
    _clear_body(document)
    _configure_page(document)
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph._element.getparent().remove(paragraph._element)
    table = header.add_table(rows=1, cols=2, width=Inches(6.65))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.95)
    table.columns[1].width = Inches(1.70)
    left, right = table.rows[0].cells
    for cell in (left, right):
        _set_cell_margins(cell, top=30, bottom=55, start=40, end=40)
    left_p = left.paragraphs[0]
    left_p.text = "DOCUMENT TITLE"
    left_p.runs[0].font.name = theme["font"]
    left_p.runs[0].font.size = Pt(8.3)
    left_p.runs[0].font.bold = True
    left_p.runs[0].font.color.rgb = _rgb(theme["primary"])
    _tag_paragraph(left_p, "AGDOC.FIELD.HEADER.TITLE", "Running title")
    right_p = right.paragraphs[0]
    right_p.text = theme["header_label"]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.runs[0].font.name = theme["font"]
    right_p.runs[0].font.size = Pt(7.2)
    right_p.runs[0].font.bold = True
    right_p.runs[0].font.color.rgb = _rgb(theme["accent"])
    _set_table_borders(table, theme["warm"], 8, top=False, bottom=True, inside=False)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def create_footer(path: Path, style_source: Path, theme_name: str) -> None:
    theme = THEMES[theme_name]
    document = Document(str(style_source))
    _set_document_properties(
        document,
        title=f"{theme_name.title()} folio footer",
        subject="Reusable Word-native footer donor",
    )
    _clear_body(document)
    _configure_page(document)
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph._element.getparent().remove(paragraph._element)
    table = footer.add_table(rows=1, cols=3, width=Inches(6.65))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (3.35, 1.45, 1.85)
    for column, width in zip(table.columns, widths):
        column.width = Inches(width)
    cells = table.rows[0].cells
    for cell in cells:
        _set_cell_margins(cell, top=55, bottom=15, start=40, end=40)
    collection = cells[0].paragraphs[0]
    collection.text = "COLLECTION"
    version = cells[1].paragraphs[0]
    version.text = "VERSION"
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page = cells[2].paragraphs[0]
    page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page.add_run("PAGE ")
    append_complex_field(page, "PAGE", "1")
    page.add_run(" / ")
    append_complex_field(page, "NUMPAGES", "1")
    for paragraph in (collection, version, page):
        for run in paragraph.runs:
            run.font.name = theme["font"]
            run.font.size = Pt(7.2)
            run.font.bold = True
            run.font.color.rgb = _rgb(theme["muted"])
    _tag_paragraph(collection, "AGDOC.FIELD.FOOTER.COLLECTION", "Collection name")
    _tag_paragraph(version, "AGDOC.FIELD.FOOTER.VERSION", "Document version")
    _set_table_borders(table, theme["warm"], 6, top=True, bottom=False, inside=False)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _font(size: int, *, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/seguisb.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.is_file():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def create_study_map(path: Path) -> None:
    image = Image.new("RGB", (1800, 760), "#F5F8F7")
    draw = ImageDraw.Draw(image)
    title = _font(54, bold=True)
    label = _font(28, bold=True)
    small = _font(23)
    draw.text((90, 58), "A short observation route", fill="#17324D", font=title)
    draw.text((92, 130), "Four fixed points · repeated walks · one consistent note format", fill="#62737B", font=small)
    points = [(190, 450), (600, 350), (1030, 475), (1530, 330)]
    draw.line(points, fill="#1E7A78", width=18, joint="curve")
    names = [
        ("01", "Entrance edge", "high movement"),
        ("02", "Shaded bench", "longer pauses"),
        ("03", "Low wall", "outside fast path"),
        ("04", "Open seats", "direct sun"),
    ]
    for index, ((x, y), (number, name, note)) in enumerate(zip(points, names)):
        radius = 54
        draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill="#17324D", outline="#D8A84E", width=8)
        box = draw.textbbox((0, 0), number, font=label)
        draw.text((x - (box[2] - box[0]) / 2, y - (box[3] - box[1]) / 2 - 4), number, fill="white", font=label)
        anchor_y = y + 82 if index % 2 == 0 else y - 142
        draw.text((x - 100, anchor_y), name, fill="#17324D", font=label)
        draw.text((x - 100, anchor_y + 38), note, fill="#62737B", font=small)
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, format="PNG", optimize=True)


def create_session_map(path: Path) -> None:
    image = Image.new("RGB", (1800, 720), "#FFFDF9")
    draw = ImageDraw.Draw(image)
    title = _font(50, bold=True)
    label = _font(32, bold=True)
    small = _font(24)
    draw.text((95, 58), "A decision people can repeat tomorrow", fill="#3D2B56", font=title)
    stages = [
        ("1", "REVEAL", "Put facts, tensions,\nand conditions in view", "#3D2B56"),
        ("2", "SHAPE", "Compare real options\nthrough one shared lens", "#D75D45"),
        ("3", "COMMIT", "Name the choice, owner,\nand review trigger", "#B68A24"),
    ]
    xs = [120, 675, 1230]
    for index, (x, (number, name, note, color)) in enumerate(zip(xs, stages)):
        draw.rounded_rectangle((x, 220, x + 450, 555), radius=36, fill=color)
        draw.ellipse((x + 34, 255, x + 110, 331), fill="#FFFDF9")
        draw.text((x + 58, 268), number, fill=color, font=label)
        draw.text((x + 42, 365), name, fill="#FFFDF9", font=label)
        draw.multiline_text((x + 42, 420), note, fill="#FFFDF9", font=small, spacing=10)
        if index < 2:
            ax = x + 470
            draw.line((ax, 385, ax + 150, 385), fill="#EFCB68", width=16)
            draw.polygon([(ax + 150, 385), (ax + 112, 360), (ax + 112, 410)], fill="#EFCB68")
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, format="PNG", optimize=True)


def _finish_workbook(workbook: Workbook, path: Path) -> None:
    workbook.properties.creator = "Agentic Word Documents contributors"
    workbook.properties.subject = "Synthetic template data"
    workbook.properties.description = "Generated by tools/build_template_assets.py"
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(path)


def create_observations(path: Path) -> None:
    workbook = Workbook()
    workbook.properties.title = "Synthetic field observations"
    sheet = workbook.active
    sheet.title = "Observations"
    rows = [
        ["Point", "Time", "Shade", "People paused", "Typical pause", "Recorder note"],
        ["01 · Entrance edge", "09:00", "Partial", 3, "Under 2 min", "Mostly waiting for another person"],
        ["02 · Shaded bench", "09:15", "Full", 7, "5–12 min", "Several people stayed after checking a phone"],
        ["03 · Low wall", "09:30", "Partial", 5, "3–7 min", "Used as an informal seat outside the fast path"],
        ["04 · Open seats", "09:45", "None", 1, "Under 2 min", "Seat used briefly; direct sun on the surface"],
        ["01 · Entrance edge", "10:00", "None", 4, "Under 2 min", "High movement, few sustained pauses"],
        ["02 · Shaded bench", "10:15", "Full", 8, "5–12 min", "Two small groups shared the bench area"],
        ["03 · Low wall", "10:30", "Partial", 6, "3–7 min", "People retained a clear view of the path"],
        ["04 · Open seats", "10:45", "None", 2, "Under 2 min", "Short stops only"],
    ]
    for row in rows:
        sheet.append(row)
    table = Table(displayName="ObservationsData", ref=f"A1:F{len(rows)}")
    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True, showFirstColumn=False)
    sheet.add_table(table)
    sheet.freeze_panes = "A2"
    widths = {"A": 23, "B": 11, "C": 12, "D": 16, "E": 16, "F": 52}
    for column, width in widths.items():
        sheet.column_dimensions[column].width = width
    for cell in sheet[1]:
        cell.fill = PatternFill("solid", fgColor="17324D")
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(vertical="center")
    _finish_workbook(workbook, path)


def create_run_of_show(path: Path) -> None:
    workbook = Workbook()
    workbook.properties.title = "Sample workshop run of show"
    sheet = workbook.active
    sheet.title = "Run of show"
    rows = [
        ["Minutes", "Movement", "Facilitator action", "Visible output"],
        ["0–10", "Frame", "State the decision boundary and what the room can influence", "One decision sentence"],
        ["10–30", "Reveal", "Collect facts, concerns, and conditions without debate", "Clustered evidence wall"],
        ["30–55", "Shape", "Name options and compare them through agreed criteria", "Comparable option cards"],
        ["55–75", "Commit", "Select, test the wording, and name an owner", "Decision record"],
        ["75–90", "Close", "Repeat the choice, list open work, confirm the review trigger", "Shared next step"],
    ]
    for row in rows:
        sheet.append(row)
    table = Table(displayName="RunOfShow", ref=f"A1:D{len(rows)}")
    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium4", showRowStripes=True, showFirstColumn=False)
    sheet.add_table(table)
    sheet.freeze_panes = "A2"
    widths = {"A": 12, "B": 15, "C": 58, "D": 31}
    for column, width in widths.items():
        sheet.column_dimensions[column].width = width
    for cell in sheet[1]:
        cell.fill = PatternFill("solid", fgColor="3D2B56")
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(vertical="center")
    _finish_workbook(workbook, path)


def create_decision_canvas(path: Path, style_source: Path) -> None:
    theme = THEMES["storyboard"]
    document = Document(str(style_source))
    _set_document_properties(
        document,
        title="Decision canvas",
        subject="Editable worksheet fragment",
    )
    _clear_body(document)
    _configure_page(document)
    heading = document.add_paragraph("Decision canvas")
    heading.style = "Heading 2"
    intro = document.add_paragraph("Complete the four boxes before the session closes. Use short, observable language.")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for column in table.columns:
        column.width = Cm(8.1)
    prompts = [
        ("THE CHOICE", "We decided…\n\n\n"),
        ("WHY THIS OPTION", "It best meets…\n\n\n"),
        ("OWNER + NEXT ACTION", "Name / action / due point…\n\n\n"),
        ("REVIEW TRIGGER", "We will revisit this when…\n\n\n"),
    ]
    for cell, (label, prompt) in zip([cell for row in table.rows for cell in row.cells], prompts):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_cell_fill(cell, theme["paper"])
        _set_cell_margins(cell, top=150, bottom=150, start=150, end=150)
        paragraph = cell.paragraphs[0]
        paragraph.text = label
        paragraph.runs[0].font.name = theme["font"]
        paragraph.runs[0].font.size = Pt(9)
        paragraph.runs[0].font.bold = True
        paragraph.runs[0].font.color.rgb = _rgb(theme["accent"])
        body = cell.add_paragraph(prompt)
        body.runs[0].font.name = theme["font"]
        body.runs[0].font.size = Pt(10)
        body.runs[0].font.color.rgb = _rgb(theme["muted"])
    for row in table.rows:
        row.height = Cm(4.0)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _set_table_borders(table, theme["primary"], 8, top=True, bottom=True, inside=True)
    elements = [child for child in list(document._element.body) if child.tag != qn("w:sectPr")]
    wrap_elements("AGDOC.BODY.DECISION_CANVAS", elements, alias="Printable decision canvas")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def main() -> None:
    for theme_name in THEMES:
        donor_root = ROOT / "kits" / theme_name / "donors"
        style = donor_root / "Style_Gallery.docx"
        create_style_donor(style, theme_name)
        create_header(donor_root / "Running_Header.docx", style, theme_name)
        create_footer(donor_root / "Folio_Footer.docx", style, theme_name)

    studio_style = ROOT / "kits" / "studio" / "donors" / "Style_Gallery.docx"
    storyboard_style = ROOT / "kits" / "storyboard" / "donors" / "Style_Gallery.docx"
    create_plain_shell(ROOT / "profiles" / "plain" / "shell.docx", studio_style)
    create_shell(ROOT / "profiles" / "report" / "shell.docx", studio_style, "AGDOC.BODY.REPORT")
    create_shell(ROOT / "profiles" / "handbook" / "shell.docx", storyboard_style, "AGDOC.BODY.HANDBOOK")

    field_root = ROOT / "projects" / "field-study" / "documents" / "shade-study"
    create_cover(field_root / "presentation" / "cover.docx", studio_style, "studio")
    create_study_map(field_root / "content" / "study-map.png")
    create_observations(field_root / "content" / "observations.xlsx")

    workshop_root = ROOT / "projects" / "workshop-kit" / "documents" / "clear-decisions"
    create_cover(workshop_root / "presentation" / "cover.docx", storyboard_style, "storyboard")
    create_session_map(workshop_root / "content" / "session-map.png")
    create_run_of_show(workshop_root / "content" / "run-of-show.xlsx")
    create_decision_canvas(workshop_root / "content" / "decision-canvas.docx", storyboard_style)

    print("Generated style donors, shells, covers, figures, workbooks, and Word fragments.")


if __name__ == "__main__":
    main()
