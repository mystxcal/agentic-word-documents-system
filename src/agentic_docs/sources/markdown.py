from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urlparse

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

from ..diagnostics import DiagnosticBag
from ..errors import PackageError
from ..word.ooxml import make_sdt
from ..word.package import build_native_table
from .markdown_ast import Block, Inline, MarkdownDocument, parse_blocks, parse_document, parse_inlines


def _safe_tag_part(value: str, width: int = 18) -> str:
    result = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("._-")
    return (result or "item")[:width]


def markdown_slot_tag(component_id: str, slot_name: str) -> str:
    digest = hashlib.sha256(f"{component_id}\0{slot_name}".encode("utf-8")).hexdigest()[:10]
    return f"AGDOC.MD.SLOT.{_safe_tag_part(component_id)}.{_safe_tag_part(slot_name)}.{digest}"


def markdown_block_tag(component_id: str, block_id: str) -> str:
    digest = hashlib.sha256(f"{component_id}\0{block_id}".encode("utf-8")).hexdigest()[:10]
    return f"AGDOC.MD.{_safe_tag_part(component_id)}.{_safe_tag_part(block_id)}.{digest}"


def markdown_component_tag(component_id: str) -> str:
    digest = hashlib.sha256(f"component\0{component_id}".encode("utf-8")).hexdigest()[:10]
    return f"AGDOC.MD.COMPONENT.{_safe_tag_part(component_id)}.{digest}"


def _style(document: Document, requested: str, fallback: str, diagnostics: DiagnosticBag, location: str) -> str:
    names = {style.name for style in document.styles}
    if requested in names:
        return requested
    if fallback in names:
        diagnostics.warn(
            "SEMANTIC_STYLE_FALLBACK",
            f"Requested Word style {requested!r} is unavailable; using {fallback!r}",
            location=location,
        )
        return fallback
    diagnostics.warn(
        "SEMANTIC_STYLE_UNAVAILABLE",
        f"Requested Word style {requested!r} and fallback {fallback!r} are unavailable; using Normal",
        location=location,
    )
    return "Normal"


def _clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _detached_paragraph(document: Document, style_name: str | None = None):
    paragraph = document.add_paragraph(style=style_name)
    document._element.body.remove(paragraph._p)
    return paragraph


def _append_document_element(document: Document, element) -> None:
    body = document._element.body
    section = body.find(qn("w:sectPr"))
    if section is None:
        body.append(element)
    else:
        body.insert(body.index(section), element)


def _new_list_num_id(document: Document, *, ordered: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "hybridMultilevel")
    abstract.append(multi)

    number_formats = ("decimal", "lowerLetter", "lowerRoman")
    bullet_marks = ("•", "◦", "▪")
    for level in range(9):
        definition = OxmlElement("w:lvl")
        definition.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        number_format = OxmlElement("w:numFmt")
        number_format.set(qn("w:val"), number_formats[level % 3] if ordered else "bullet")
        level_text = OxmlElement("w:lvlText")
        level_text.set(
            qn("w:val"),
            f"%{level + 1}." if ordered else bullet_marks[level % 3],
        )
        suffix = OxmlElement("w:suff")
        suffix.set(qn("w:val"), "tab")
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        paragraph_properties = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(360 * (level + 1)))
        tabs.append(tab)
        indentation = OxmlElement("w:ind")
        indentation.set(qn("w:left"), str(360 * (level + 1)))
        indentation.set(qn("w:hanging"), "360")
        paragraph_properties.extend((tabs, indentation))
        definition.extend((start, number_format, level_text, suffix, justification, paragraph_properties))
        if not ordered:
            run_properties = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Arial")
            fonts.set(qn("w:hAnsi"), "Arial")
            run_properties.append(fonts)
            definition.append(run_properties)
        abstract.append(definition)

    first_num = next(
        (index for index, element in enumerate(numbering) if element.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num, abstract)

    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    reference = OxmlElement("w:abstractNumId")
    reference.set(qn("w:val"), str(abstract_id))
    number.append(reference)
    numbering.append(number)
    return num_id


def _apply_num(paragraph, num_id: int | None, level: int) -> None:
    if num_id is None:
        return
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(8, max(0, level))))
    value = OxmlElement("w:numId")
    value.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, value))
    p_pr.append(num_pr)


def _inline_text(nodes: Iterable[Inline]) -> str:
    result = []
    for node in nodes:
        if node.kind == "hard_break":
            result.append(" ")
        elif node.children:
            result.append(_inline_text(node.children))
        else:
            result.append(node.text)
    return "".join(result)


def _valid_link(target: str) -> bool:
    parsed = urlparse(target)
    return parsed.scheme in {"http", "https", "mailto"} and bool(parsed.path or parsed.netloc)


def _render_inlines(paragraph, nodes: Iterable[Inline], inherited: frozenset[str] = frozenset()) -> None:
    for node in nodes:
        if node.kind == "hard_break":
            paragraph.add_run().add_break()
            continue
        if node.kind == "link":
            label = _inline_text(node.children)
            if not node.target or not _valid_link(node.target):
                paragraph.add_run(label + (f" ({node.target})" if node.target else ""))
                continue
            relationship = paragraph.part.relate_to(
                node.target,
                RELATIONSHIP_TYPE.HYPERLINK,
                is_external=True,
            )
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), relationship)
            run = OxmlElement("w:r")
            properties = OxmlElement("w:rPr")
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "0563C1")
            underline = OxmlElement("w:u")
            underline.set(qn("w:val"), "single")
            properties.extend((color, underline))
            text = OxmlElement("w:t")
            text.text = label
            run.extend((properties, text))
            hyperlink.append(run)
            paragraph._p.append(hyperlink)
            continue
        formats = inherited
        if node.kind in {"strong", "emphasis", "strike"}:
            formats = frozenset({*formats, node.kind})
            _render_inlines(paragraph, node.children, formats)
            continue
        value = node.text
        run = paragraph.add_run(value)
        run.bold = "strong" in formats
        run.italic = "emphasis" in formats
        run.font.strike = "strike" in formats
        if node.kind == "code":
            run.font.name = "Consolas"


def _table_dataset(block: Block) -> dict[str, Any]:
    headers = list(block.value["headers"])
    alignments = list(block.value["alignments"])
    columns = [
        {
            "source": f"column_{index + 1}",
            "heading": heading,
            "align": alignments[index],
            "header_align": "center",
        }
        for index, heading in enumerate(headers)
    ]
    records = [
        {
            "id": None,
            "group": None,
            "values": list(row),
        }
        for row in block.value["rows"]
    ]
    return {"columns": columns, "records": records}


def _render_block(
    document: Document,
    block: Block,
    *,
    component_id: str,
    styles: dict[str, str],
    table_styles: dict[str, dict[str, Any]],
    table_role: str,
    total_width_twips: int,
    diagnostics: DiagnosticBag,
    location: str,
    callout_role: str | None = None,
) -> list:
    if block.kind == "heading":
        level = min(int(block.value["level"]), 6)
        fallback = f"Heading {level}"
        paragraph = _detached_paragraph(
            document,
            _style(document, styles.get(f"heading_{level}", fallback), fallback, diagnostics, location),
        )
        _render_inlines(paragraph, parse_inlines(str(block.value["text"])))
        return [paragraph._p]
    if block.kind in {"paragraph", "quote"}:
        role = callout_role or ("quote" if block.kind == "quote" else "body")
        fallback = "Quote" if block.kind == "quote" else "Normal"
        paragraph = _detached_paragraph(
            document,
            _style(document, styles.get(role, styles.get("body", fallback)), fallback, diagnostics, location),
        )
        _render_inlines(paragraph, parse_inlines(str(block.value)))
        return [paragraph._p]
    if block.kind == "list":
        style_names = {
            True: _style(
                document,
                styles.get("numbered_step", "List Number"),
                "List Number",
                diagnostics,
                location,
            ),
            False: _style(
                document,
                styles.get("bullet", "List Bullet"),
                "List Bullet",
                diagnostics,
                location,
            ),
        }
        used_types = {bool(item["ordered"]) for item in block.value["items"]}
        num_ids = {ordered: _new_list_num_id(document, ordered=ordered) for ordered in used_types}
        elements = []
        for item in block.value["items"]:
            ordered = bool(item["ordered"])
            paragraph = _detached_paragraph(document, style_names[ordered])
            _apply_num(paragraph, num_ids[ordered], int(item["level"]))
            _render_inlines(paragraph, parse_inlines(str(item["text"])))
            elements.append(paragraph._p)
        return elements
    if block.kind == "code":
        style_name = _style(
            document,
            styles.get("command", "Normal"),
            "Normal",
            diagnostics,
            location,
        )
        elements = []
        for line in str(block.value["text"]).split("\n"):
            paragraph = _detached_paragraph(document, style_name)
            run = paragraph.add_run(line)
            run.font.name = "Consolas"
            elements.append(paragraph._p)
        return elements
    if block.kind == "page_break":
        paragraph = _detached_paragraph(document)
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        return [paragraph._p]
    if block.kind == "table":
        if table_role not in table_styles:
            raise PackageError(f"Markdown table selects undefined table style role {table_role!r}")
        return [
            build_native_table(
                _table_dataset(block),
                table_styles[table_role],
                total_width_twips,
                "No rows supplied.",
            )
        ]
    if block.kind == "insert":
        slot_name = str(block.value["slot"])
        placeholder = OxmlElement("w:p")
        return [
            make_sdt(
                markdown_slot_tag(component_id, slot_name),
                [placeholder],
                alias=f"Markdown+ insertion: {slot_name}",
                sdt_id=int(hashlib.sha256(f"slot:{component_id}:{slot_name}".encode()).hexdigest()[:7], 16),
            )
        ]
    if block.kind == "callout":
        role = str(block.value["role"])
        effective_role = role if role in styles else "note"
        elements = []
        for child in block.value["blocks"]:
            elements.extend(
                _render_block(
                    document,
                    child,
                    component_id=component_id,
                    styles=styles,
                    table_styles=table_styles,
                    table_role=table_role,
                    total_width_twips=total_width_twips,
                    diagnostics=diagnostics,
                    location=location,
                    callout_role=effective_role,
                )
            )
        if not elements:
            raise PackageError(f"Markdown+ callout at {location}:{block.start_line} is empty")
        return elements
    raise PackageError(f"Unsupported parsed Markdown block {block.kind!r}")


def compile_markdown(
    source_path: Path,
    output_path: Path,
    *,
    component_id: str,
    style_source: Path,
    semantic_styles: dict[str, str],
    table_styles: dict[str, dict[str, Any]],
    declared_slots: dict[str, list[str]],
    options: dict[str, Any],
    total_width_twips: int,
    diagnostics: DiagnosticBag,
) -> dict[str, Any]:
    """Compile canonical Markdown+ prose into one style-native Word fragment."""

    source_path = Path(source_path).resolve()
    output_path = Path(output_path).resolve()
    text = source_path.read_text(encoding="utf-8-sig")
    strict = bool(options.get("strict", True))
    parsed: MarkdownDocument = parse_document(
        text,
        diagnostics,
        str(source_path),
        strict=strict,
    )
    declared = list(declared_slots)
    unknown = sorted(set(parsed.slot_names) - set(declared))
    missing = sorted(set(declared) - set(parsed.slot_names))
    if unknown:
        raise PackageError(
            f"Markdown+ source {source_path} inserts undeclared component slot(s): {', '.join(unknown)}"
        )
    if missing:
        raise PackageError(
            f"Markdown+ source {source_path} does not place declared component slot(s): {', '.join(missing)}"
        )

    styles = dict(semantic_styles)
    style_overrides = options.get("style_roles") or {}
    if not isinstance(style_overrides, dict) or not all(
        isinstance(key, str) and isinstance(value, str) for key, value in style_overrides.items()
    ):
        raise PackageError("Markdown options.style_roles must map semantic role names to Word style names")
    styles.update(style_overrides)
    table_role = str(options.get("table_style_role", "technical"))

    document = Document(str(style_source))
    _clear_body(document)
    source_map = []
    compiled_blocks = []
    for block in parsed.blocks:
        elements = _render_block(
            document,
            block,
            component_id=component_id,
            styles=styles,
            table_styles=table_styles,
            table_role=table_role,
            total_width_twips=total_width_twips,
            diagnostics=diagnostics,
            location=str(source_path),
        )
        tag = markdown_block_tag(component_id, block.block_id or block.kind)
        alias = f"{source_path.name}:{block.start_line}-{block.end_line}"
        sdt_id = int(hashlib.sha256(tag.encode()).hexdigest()[:7], 16)
        if block.kind == "heading" and len(elements) == 1:
            # A heading must remain a direct paragraph in the document flow.
            # Word creates TOC bookmarks around heading paragraphs; wrapping
            # that paragraph in a block content control makes PAGEREF results
            # unreliable during PDF export. Keep source ownership at run level
            # inside the paragraph instead, so the heading is still tagged and
            # editable without compromising native TOC behavior.
            paragraph = elements[0]
            paragraph_properties = paragraph.find(qn("w:pPr"))
            content = [child for child in list(paragraph) if child is not paragraph_properties]
            for child in content:
                paragraph.remove(child)
            paragraph.append(make_sdt(tag, content, alias=alias, sdt_id=sdt_id))
            compiled_blocks.append(paragraph)
        else:
            compiled_blocks.append(make_sdt(tag, elements, alias=alias, sdt_id=sdt_id))
        source_map.append(
            {
                "block_id": block.block_id,
                "kind": block.kind,
                "source_lines": [block.start_line, block.end_line],
                "word_tag": tag,
            }
        )

    root_tag = markdown_component_tag(component_id)
    root_wrapper = make_sdt(
        root_tag,
        compiled_blocks,
        alias=f"Markdown+ component: {component_id}",
        sdt_id=int(hashlib.sha256(root_tag.encode()).hexdigest()[:7], 16),
    )
    _append_document_element(document, root_wrapper)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    temporary = output_path.with_suffix(output_path.suffix + ".tmp")
    document.save(temporary)
    temporary.replace(output_path)
    diagnostics.info(
        "MARKDOWN_PLUS_COMPILED",
        f"Compiled {len(parsed.blocks)} Markdown+ block(s) and {len(parsed.slot_names)} insertion slot(s) for {component_id!r}",
        location=str(source_path),
    )
    return {
        "schema": "agentic-markdown-compile/v1",
        "source": str(source_path),
        "output": str(output_path),
        "component": component_id,
        "component_tag": root_tag,
        "block_count": len(parsed.blocks),
        "slots": parsed.slot_names,
        "source_map": source_map,
        "table_style_role": table_role,
    }


__all__ = [
    "Block",
    "compile_markdown",
    "markdown_block_tag",
    "markdown_component_tag",
    "markdown_slot_tag",
    "parse_blocks",
    "parse_document",
    "parse_inlines",
]
