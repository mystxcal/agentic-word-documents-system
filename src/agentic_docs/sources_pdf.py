from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from .errors import PackageError
from .rendering import render_pdf_pages
from .word.ooxml import qn, wrap_elements


def expand_page_selection(value: Any) -> list[int]:
    """Expand integers, [start, end] ranges, or {start, end} ranges in order."""

    if not isinstance(value, list) or not value:
        raise PackageError("pdf_pages requires a non-empty options.pages list")
    pages: list[int] = []
    for item in value:
        if isinstance(item, int) and not isinstance(item, bool):
            start = end = item
        elif isinstance(item, (list, tuple)) and len(item) == 2:
            start, end = item
        elif isinstance(item, dict) and set(item) >= {"start", "end"}:
            start, end = item["start"], item["end"]
        elif hasattr(item, "start") and hasattr(item, "end"):
            start, end = item.start, item.end
        else:
            raise PackageError(f"Unsupported PDF page selector: {item!r}")
        if (
            not isinstance(start, int)
            or isinstance(start, bool)
            or not isinstance(end, int)
            or isinstance(end, bool)
            or start < 1
            or end < start
        ):
            raise PackageError(f"Invalid PDF page range: {start!r}-{end!r}")
        pages.extend(range(start, end + 1))
    return pages


def _empty_document(style_source: Path) -> Document:
    document = Document(str(style_source))
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    return document


def compile_pdf_pages(
    source: Path,
    output: Path,
    *,
    style_source: Path,
    options: dict[str, Any],
    title: str | None,
    caption: str | bool | None,
    alt_text: str | None,
    work_directory: Path,
    available_width_inches: float,
    source_tag: str | None = None,
) -> dict[str, Any]:
    pages = expand_page_selection(options.get("pages"))
    dpi = int(options.get("dpi", 150))
    if dpi < 72 or dpi > 600:
        raise PackageError("pdf_pages options.dpi must be between 72 and 600")
    rendered = render_pdf_pages(source, work_directory, pages, dpi=dpi)
    document = _empty_document(style_source)
    configured_width = options.get("image_width_inches")
    width = float(configured_width if configured_width is not None else available_width_inches)
    if width <= 0 or width > available_width_inches + 0.05:
        raise PackageError(
            f"pdf_pages image width {width} exceeds the available document width {available_width_inches:.2f} inches"
        )
    alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(str(options.get("alignment", "center")).lower())
    if alignment is None:
        raise PackageError("pdf_pages alignment must be left, center, or right")

    if title:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(title)
        run.bold = True
    note = options.get("note")
    if note:
        paragraph = document.add_paragraph(str(note))
        paragraph.paragraph_format.keep_with_next = True

    total = len(rendered["pages"])
    for index, item in enumerate(rendered["pages"], 1):
        if caption:
            template = (
                "{source_name} | Source PDF page {page}"
                if caption is True
                else str(caption)
            )
            try:
                value = template.format(
                    source_name=source.name,
                    source_stem=source.stem,
                    page=item["source_page"],
                    index=index,
                    total=total,
                )
            except (KeyError, ValueError) as exc:
                raise PackageError(f"Invalid pdf_pages caption template: {exc}") from exc
            paragraph = document.add_paragraph(value)
            if index > 1 and options.get("page_break_between", True):
                paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.keep_with_next = True
            paragraph.alignment = alignment
        picture = document.add_paragraph()
        if not caption and index > 1 and options.get("page_break_between", True):
            picture.paragraph_format.page_break_before = True
        picture.alignment = alignment
        shape = picture.add_run().add_picture(item["image"], width=Inches(width))
        description = alt_text or f"{source.name}, source PDF page {item['source_page']}"
        shape._inline.docPr.set("descr", description)

    if source_tag:
        body_elements = [child for child in list(document._element.body) if child.tag != qn("w:sectPr")]
        wrap_elements(source_tag, body_elements, alias="Generated selected PDF pages")
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return {
        "source": str(source),
        "pages": pages,
        "dpi": dpi,
        "page_count": len(pages),
        "output": str(output),
    }
