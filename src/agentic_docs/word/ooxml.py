"""Small, explicit OOXML helpers for the Agentic Word Documents system.

The module intentionally uses public WordprocessingML structures rather than a
private templating engine.  It is shared by bootstrap, scoped refresh, and QA
tools so that tags, fields, and package metadata are handled consistently.
"""

from __future__ import annotations

import copy
import hashlib
import json
import re
import zipfile
from pathlib import Path
from typing import Iterable

from lxml import etree
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
VT_NS = "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"

NS = {
    "w": W_NS,
    "r": R_NS,
    "cp": CP_NS,
    "vt": VT_NS,
    "pr": PKG_REL_NS,
    "ct": CT_NS,
}


def xpath(element, expression: str, **variables):
    """Run one namespace-stable XPath against python-docx or plain lxml nodes."""
    return etree.XPath(expression, namespaces=NS)(element, **variables)


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest().upper()


def canonical_json_hash(value: object) -> str:
    payload = json.dumps(value, sort_keys=True, separators=(",", ":"), ensure_ascii=False)
    return hashlib.sha256(payload.encode("utf-8")).hexdigest().upper()


def sdt_tag(element) -> str | None:
    if element.tag != qn("w:sdt"):
        return None
    tag = element.find("./" + qn("w:sdtPr") + "/" + qn("w:tag"))
    return tag.get(qn("w:val")) if tag is not None else None


def make_sdt(tag: str, elements: Iterable, *, alias: str | None = None, sdt_id: int = 1):
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    alias_el = OxmlElement("w:alias")
    alias_el.set(qn("w:val"), alias or tag)
    tag_el = OxmlElement("w:tag")
    tag_el.set(qn("w:val"), tag)
    id_el = OxmlElement("w:id")
    id_el.set(qn("w:val"), str(sdt_id))
    sdt_pr.extend([alias_el, tag_el, id_el])
    content = OxmlElement("w:sdtContent")
    for element in elements:
        content.append(element)
    sdt.extend([sdt_pr, content])
    return sdt


def next_sdt_id(root) -> int:
    ids: list[int] = []
    for node in xpath(root, ".//w:sdtPr/w:id"):
        raw = node.get(qn("w:val"))
        try:
            ids.append(int(raw))
        except (TypeError, ValueError):
            continue
    return max(ids, default=0) + 1


def wrap_elements(tag: str, elements: list, *, alias: str | None = None):
    if not elements:
        raise ValueError(f"Cannot create empty content control {tag}")
    parent = elements[0].getparent()
    if parent is None or any(el.getparent() is not parent for el in elements):
        raise ValueError(f"Elements for {tag} do not share one parent")
    indices = [parent.index(el) for el in elements]
    expected = list(range(indices[0], indices[0] + len(indices)))
    if indices != expected:
        raise ValueError(f"Elements for {tag} are not contiguous")
    sdt = make_sdt(tag, [copy.deepcopy(el) for el in elements], alias=alias, sdt_id=next_sdt_id(parent.getroottree()))
    parent.insert(indices[0], sdt)
    for element in elements:
        parent.remove(element)
    return sdt


def find_sdts(root, tag: str):
    return [
        sdt
        for sdt in xpath(root, ".//w:sdt")
        if sdt_tag(sdt) == tag
    ]


def content_control_inventory(root) -> list[dict[str, object]]:
    inventory: list[dict[str, object]] = []
    for index, sdt in enumerate(xpath(root, ".//w:sdt"), start=1):
        tag = sdt_tag(sdt)
        text = "".join(xpath(sdt, ".//w:t/text()")).strip()
        inventory.append(
            {
                "index": index,
                "tag": tag,
                "text_preview": re.sub(r"\s+", " ", text)[:160],
                "tables": len(xpath(sdt, ".//w:tbl")),
                "drawings": len(xpath(sdt, ".//w:drawing|.//w:pict")),
            }
        )
    return inventory


def append_complex_field(paragraph, instruction: str, cached_text: str = "") -> None:
    parent = paragraph._p
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    instruction_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = f" {instruction} "
    instruction_run.append(instr)
    separator_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separator_run.append(separate)
    text_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = cached_text
    text_run.append(text)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    parent.extend([begin_run, instruction_run, separator_run, text_run, end_run])


def replace_with_caption_field(
    paragraph,
    *,
    label: str,
    number: int,
    description: str,
    bookmark_name: str,
    bookmark_id: int,
) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    paragraph.add_run(f"{label} ")
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    p.append(start)
    append_complex_field(paragraph, f"SEQ {label} \\* ARABIC", str(number))
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p.append(end)
    paragraph.add_run(f" — {description}")


def add_hidden_identifier(cell, identifier: str) -> None:
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(f"{{#{identifier}}}")
    vanish = OxmlElement("w:vanish")
    run._r.get_or_add_rPr().append(vanish)


def table_text(table_element) -> list[list[str]]:
    rows: list[list[str]] = []
    for tr in xpath(table_element, "./w:tr"):
        row: list[str] = []
        for tc in xpath(tr, "./w:tc"):
            text = "".join(xpath(tc, ".//w:t/text()"))
            row.append(re.sub(r"\s+", " ", text).strip())
        rows.append(row)
    return rows


def package_parts(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as zf:
        return {
            name: hashlib.sha256(zf.read(name)).hexdigest().upper()
            for name in sorted(zf.namelist())
        }


def set_custom_properties(path: Path, properties: dict[str, object]) -> None:
    """Add/replace scalar custom properties without disturbing other parts."""

    with zipfile.ZipFile(path) as zf:
        payloads = {name: zf.read(name) for name in zf.namelist()}

    custom_name = "docProps/custom.xml"
    if custom_name in payloads:
        root = etree.fromstring(payloads[custom_name])
    else:
        root = etree.Element(f"{{{CP_NS}}}Properties", nsmap={None: CP_NS, "vt": VT_NS})

    existing = {
        node.get("name"): node
        for node in root.findall(f"{{{CP_NS}}}property")
        if node.get("name")
    }
    next_pid = max([int(node.get("pid", "1")) for node in existing.values()] + [1]) + 1
    fmtid = "{D5CDD505-2E9C-101B-9397-08002B2CF9AE}"
    for name, value in properties.items():
        node = existing.get(name)
        if node is None:
            node = etree.SubElement(
                root,
                f"{{{CP_NS}}}property",
                fmtid=fmtid,
                pid=str(next_pid),
                name=name,
            )
            next_pid += 1
        for child in list(node):
            node.remove(child)
        if isinstance(value, bool):
            child = etree.SubElement(node, f"{{{VT_NS}}}bool")
            child.text = "true" if value else "false"
        elif isinstance(value, int):
            child = etree.SubElement(node, f"{{{VT_NS}}}i4")
            child.text = str(value)
        else:
            child = etree.SubElement(node, f"{{{VT_NS}}}lpwstr")
            child.text = "" if value is None else str(value)
    payloads[custom_name] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")

    rels_name = "_rels/.rels"
    rels = etree.fromstring(payloads[rels_name])
    custom_rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties"
    if not rels.xpath("./pr:Relationship[@Type=$t]", namespaces=NS, t=custom_rel_type):
        ids = {
            int(match.group(1))
            for rid in rels.xpath("./pr:Relationship/@Id", namespaces=NS)
            if (match := re.fullmatch(r"rId(\d+)", rid))
        }
        rid = 1
        while rid in ids:
            rid += 1
        etree.SubElement(
            rels,
            f"{{{PKG_REL_NS}}}Relationship",
            Id=f"rId{rid}",
            Type=custom_rel_type,
            Target="docProps/custom.xml",
        )
        payloads[rels_name] = etree.tostring(rels, xml_declaration=True, encoding="UTF-8", standalone="yes")

    types_name = "[Content_Types].xml"
    types = etree.fromstring(payloads[types_name])
    if not types.xpath("./ct:Override[@PartName='/docProps/custom.xml']", namespaces=NS):
        etree.SubElement(
            types,
            f"{{{CT_NS}}}Override",
            PartName="/docProps/custom.xml",
            ContentType="application/vnd.openxmlformats-officedocument.custom-properties+xml",
        )
        payloads[types_name] = etree.tostring(types, xml_declaration=True, encoding="UTF-8", standalone="yes")

    temp = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(temp, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for name, data in payloads.items():
            zf.writestr(name, data)
    temp.replace(path)
