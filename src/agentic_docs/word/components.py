from __future__ import annotations

from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from lxml import etree

from ..diagnostics import DiagnosticBag
from ..component_cache import ComponentAdapterCache, adapter_payload
from ..errors import PackageError
from ..model import ResolvedComponent, ResolvedDocument
from ..sources.markdown import compile_markdown, markdown_component_tag
from ..sources_excel import export_dataset
from ..sources_pdf import compile_pdf_pages
from .fragments import component_wrapper, import_word_fragment
from .ooxml import qn, wrap_elements
from .package import DocxPackage, build_native_table, document_usable_width


def _preview_placeholder(
    component: ResolvedComponent,
    diagnostics: DiagnosticBag,
    *,
    action: str,
) -> etree._Element:
    declaration = component.declaration
    if action == "omit":
        diagnostics.info(
            "PREVIEW_COMPONENT_OMITTED",
            f"Omitted component {component.id!r} from the lightweight preview",
            location=str(component.source_path) if component.source_path else None,
        )
        return component_wrapper(component.id, [])

    configured = declaration.preview.label if declaration.preview else None
    label = configured or declaration.title or component.id.replace("-", " ").replace("_", " ").title()
    detail = "Deferred in this lightweight preview; included by the complete build."
    if declaration.type.value == "pdf_pages":
        selected = declaration.options.get("pages", [])
        detail = (
            f"{len(selected)} selected source-page range(s) deferred in this lightweight preview; "
            "the complete build renders and inserts them."
        )
    paragraph = etree.Element(qn("w:p"))
    paragraph_properties = etree.SubElement(paragraph, qn("w:pPr"))
    spacing = etree.SubElement(paragraph_properties, qn("w:spacing"))
    spacing.set(qn("w:before"), "120")
    spacing.set(qn("w:after"), "120")
    border = etree.SubElement(paragraph_properties, qn("w:pBdr"))
    bottom = etree.SubElement(border, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "A6B3BA")
    run = etree.SubElement(paragraph, qn("w:r"))
    run_properties = etree.SubElement(run, qn("w:rPr"))
    etree.SubElement(run_properties, qn("w:b"))
    text = etree.SubElement(run, qn("w:t"))
    text.text = f"Preview placeholder — {label}. "
    detail_run = etree.SubElement(paragraph, qn("w:r"))
    detail_text = etree.SubElement(detail_run, qn("w:t"))
    detail_text.text = detail
    diagnostics.info(
        "PREVIEW_COMPONENT_DEFERRED",
        f"Replaced component {component.id!r} with a visible lightweight-preview placeholder",
        location=str(component.source_path) if component.source_path else None,
    )
    return component_wrapper(component.id, [paragraph])


def _materialize_generated_fragment(
    *,
    resolved: ResolvedDocument,
    component: ResolvedComponent,
    target: DocxPackage,
    build_work: Path,
    cache: ComponentAdapterCache | None,
    cache_events: list[dict] | None,
    builder: Callable[[Path], object],
) -> tuple[Path, object | None, bool]:
    compiled = build_work / "components" / f"{component.id}.docx"
    width = document_usable_width(target.xml("word/document.xml"))
    payload = adapter_payload(resolved, component, available_width_twips=width) if cache else None
    if cache is not None and payload is not None:
        cached, event = cache.lookup(payload)
        if cached is not None:
            if cache_events is not None:
                cache_events.append(event)
            return cached, None, True

    compiled.parent.mkdir(parents=True, exist_ok=True)
    builder_result = builder(compiled)
    if cache is not None and payload is not None:
        stored, event = cache.store(payload, compiled)
        if cache_events is not None:
            cache_events.append(event)
        return stored, builder_result, False
    return compiled, builder_result, False


def _empty_fragment(style_source: Path) -> Document:
    document = Document(str(style_source))
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    return document


def _figure_fragment(
    source: Path,
    output: Path,
    *,
    style_source: Path,
    alignment: str,
    caption: str | bool | None,
    title: str | None,
    available_width_inches: float,
    width_inches: float | None,
    alt_text: str | None,
    source_tag: str,
) -> None:
    if source.suffix.lower() not in {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff"}:
        raise PackageError(
            f"Figure format {source.suffix!r} is not yet supported by the native raster adapter; export a reviewed PNG/JPEG snapshot"
        )
    document = _empty_fragment(style_source)
    paragraph = document.add_paragraph()
    paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[alignment]
    selected_width = width_inches if width_inches is not None else available_width_inches
    if selected_width <= 0 or selected_width > available_width_inches + 0.05:
        raise PackageError(
            f"Figure width {selected_width:.2f} exceeds the available document width "
            f"{available_width_inches:.2f} inches"
        )
    shape = paragraph.add_run().add_picture(str(source), width=Inches(selected_width))
    shape._inline.docPr.set("descr", alt_text or source.name)
    if caption:
        value = title if caption is True else str(caption)
        caption_paragraph = document.add_paragraph(value or source.stem)
        caption_paragraph.style = "Caption" if "Caption" in {style.name for style in document.styles} else "Normal"
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body_elements = [child for child in list(document._element.body) if child.tag != qn("w:sectPr")]
    wrap_elements(source_tag, body_elements, alias="Generated publication figure")
    document.save(output)


def compile_component_wrapper(
    target: DocxPackage,
    resolved: ResolvedDocument,
    component: ResolvedComponent,
    *,
    build_work: Path,
    diagnostics: DiagnosticBag,
    cache: ComponentAdapterCache | None = None,
    cache_events: list[dict] | None = None,
    preview_action: str = "include",
) -> etree._Element | None:
    declaration = component.declaration
    component_id = component.id
    if declaration.type.value in {"cover", "toc"}:
        return None
    if preview_action in {"placeholder", "omit"}:
        return _preview_placeholder(component, diagnostics, action=preview_action)
    if declaration.type.value == "page_break":
        paragraph = etree.Element(qn("w:p"))
        run = etree.SubElement(paragraph, qn("w:r"))
        br = etree.SubElement(run, qn("w:br"))
        br.set(qn("w:type"), "page")
        return component_wrapper(component_id, [paragraph])
    if component.source_path is None:
        paragraph = etree.Element(qn("w:p"))
        run = etree.SubElement(paragraph, qn("w:r"))
        text = etree.SubElement(run, qn("w:t"))
        text.text = f"[Draft component unavailable: {component_id}]"
        diagnostics.warn(
            "DRAFT_COMPONENT_PLACEHOLDER",
            f"Inserted a visible draft placeholder for unavailable optional component {component_id!r}",
        )
        return component_wrapper(component_id, [paragraph])
    suffix = component.source_path.suffix.lower()
    preserve_sections = bool(declaration.options.get("preserve_sections", False))

    if declaration.type.value == "document" and suffix == ".docx":
        return import_word_fragment(
            target,
            component.source_path,
            component_id=component_id,
            source_tag=declaration.source_tag,
            allow_untagged=declaration.allow_untagged,
            preserve_sections=preserve_sections,
            diagnostics=diagnostics,
        )
    if declaration.type.value == "document" and suffix in {".md", ".markdown"}:
        compiled, markdown_result, cache_hit = _materialize_generated_fragment(
            resolved=resolved,
            component=component,
            target=target,
            build_work=build_work,
            cache=cache,
            cache_events=cache_events,
            builder=lambda output: compile_markdown(
                component.source_path,
                output,
                component_id=component_id,
                style_source=resolved.presentation_paths["styles"],
                semantic_styles=resolved.kit.semantic_styles,
                table_styles=resolved.kit.table_styles,
                declared_slots=declaration.slots,
                options=declaration.options,
                total_width_twips=document_usable_width(target.xml("word/document.xml")),
                diagnostics=diagnostics,
            ),
        )
        if cache_hit:
            diagnostics.info(
                "COMPONENT_ADAPTER_CACHE_HIT",
                f"Reused verified Markdown+ adapter output for component {component_id!r}",
                location=str(component.source_path),
            )
        else:
            diagnostics.info(
                "MARKDOWN_PLUS_SOURCE_MAP",
                f"Markdown+ component {component_id!r} emitted {len(markdown_result['source_map'])} source-mapped block(s)",
                location=str(component.source_path),
            )
        return import_word_fragment(
            target,
            compiled,
            component_id=component_id,
            source_tag=markdown_component_tag(component_id),
            allow_untagged=False,
            preserve_sections=False,
            diagnostics=diagnostics,
        )
    if declaration.type.value in {"figure", "diagram"}:
        generated_source_tag = f"AGDOC.GENERATED.{component_id}"
        publication_source = (
            component.related_paths["rendition"]
            if declaration.type.value == "diagram"
            else component.source_path
        )
        compiled, _figure_result, cache_hit = _materialize_generated_fragment(
            resolved=resolved,
            component=component,
            target=target,
            build_work=build_work,
            cache=cache,
            cache_events=cache_events,
            builder=lambda output: _figure_fragment(
                publication_source,
                output,
                style_source=resolved.presentation_paths["styles"],
                alignment=declaration.alignment,
                caption=declaration.caption,
                title=declaration.title,
                available_width_inches=document_usable_width(target.xml("word/document.xml")) / 1440,
                width_inches=declaration.options.get("width_inches"),
                alt_text=declaration.alt_text,
                source_tag=generated_source_tag,
            ),
        )
        if cache_hit:
            diagnostics.info(
                "COMPONENT_ADAPTER_CACHE_HIT",
                f"Reused verified {declaration.type.value} publication output for component {component_id!r}",
                location=str(publication_source),
            )
        elif declaration.type.value == "diagram":
            diagnostics.info(
                "DIAGRAM_RENDITION_EMBEDDED",
                f"Embedded the reviewed rendition associated with native diagram {component_id!r}",
                location=str(publication_source),
            )
        return import_word_fragment(
            target,
            compiled,
            component_id=component_id,
            source_tag=generated_source_tag,
            allow_untagged=False,
            preserve_sections=False,
            diagnostics=diagnostics,
        )
    if declaration.type.value == "pdf_pages" and suffix == ".pdf":
        generated_source_tag = f"AGDOC.GENERATED.{component_id}"
        render_directory = build_work / "pdf-pages" / component_id
        compiled, result, cache_hit = _materialize_generated_fragment(
            resolved=resolved,
            component=component,
            target=target,
            build_work=build_work,
            cache=cache,
            cache_events=cache_events,
            builder=lambda output: compile_pdf_pages(
                component.source_path,
                output,
                style_source=resolved.presentation_paths["styles"],
                options=declaration.options,
                title=declaration.title,
                caption=declaration.caption,
                alt_text=declaration.alt_text,
                work_directory=render_directory,
                available_width_inches=document_usable_width(target.xml("word/document.xml")) / 1440,
                source_tag=generated_source_tag,
            ),
        )
        if cache_hit:
            diagnostics.info(
                "COMPONENT_ADAPTER_CACHE_HIT",
                f"Reused verified PDF-page adapter output for component {component_id!r}",
                location=str(component.source_path),
            )
        else:
            diagnostics.info(
                "PDF_PAGES_RENDERED",
                f"Rendered {result['page_count']} selected page(s) for component {component_id!r}",
                location=str(component.source_path),
            )
        return import_word_fragment(
            target,
            compiled,
            component_id=component_id,
            source_tag=generated_source_tag,
            allow_untagged=False,
            preserve_sections=False,
            diagnostics=diagnostics,
        )
    if declaration.type.value == "table":
        source = declaration.options.get("source") or {}
        locator = source.get("locator") or declaration.options.get("locator") or {}
        view = declaration.options.get("view") or {}
        formula_policy = declaration.options.get("formula_policy", "cached_values")
        dataset = export_dataset(
            component.source_path,
            locator,
            view,
            component_id=component_id,
            formula_policy=formula_policy,
        )
        if dataset["formula_evidence"]["formula_count"] and formula_policy == "cached_values":
            diagnostics.warn(
                "EXCEL_CACHED_FORMULAS_USED",
                f"Table component {component_id!r} used cached results from "
                f"{dataset['formula_evidence']['formula_count']} Excel formula cell(s)",
                location=str(component.source_path),
                hint="Use require_no_formulas or require_cached_results when the document needs a stricter formula policy.",
            )
        role = view.get("style_role", declaration.options.get("style_role", "technical"))
        if role not in resolved.kit.table_styles:
            raise PackageError(f"Table component {component_id!r} selects undefined kit table style {role!r}")
        table = build_native_table(
            dataset,
            resolved.kit.table_styles[role],
            document_usable_width(target.xml("word/document.xml")),
            view.get("empty_text", "No rows selected."),
        )
        return component_wrapper(component_id, [table])
    raise PackageError(
        f"Component {component_id!r} has unsupported type/source combination: {declaration.type}/{suffix}"
    )
